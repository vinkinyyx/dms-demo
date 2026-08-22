# -*- coding: utf-8 -*-
"""
生成《计算机软件著作权登记申请表》Word 文档
作者：DMS 项目组
"""
import os
from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = r"d:\Workspace\TRAE\DMS\软著申请材料\01_申请表"
OUT_FILE = os.path.join(OUT_DIR, "计算机软件著作权登记申请表_DMS经销商管理系统V1.0.docx")

SOFT_FULL = "DMS经销商管理系统"
SOFT_SHORT = "DMS系统"
VERSION = "V1.0"


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            tag = qn(f'w:{edge}')
            el = tcBorders.find(tag)
            if el is None:
                el = OxmlElement(f'w:{edge}')
                tcBorders.append(el)
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), '6')
            el.set(qn('w:color'), '000000')


def set_run_font(run, name=u'宋体', size=10.5, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)


def add_para(cell_or_doc, text, size=10.5, bold=False, align=None, font=u'宋体'):
    p = cell_or_doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, name=font, size=size, bold=bold)
    return p


def fill_cell(cell, text, bold=False, size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def shade_cell(cell, color_hex="D9E1F2"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    doc = Document()

    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # 标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("计算机软件著作权登记申请表")
    set_run_font(r, name=u'宋体', size=18, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("（中国版权保护中心 制）")
    set_run_font(r, size=10.5)

    doc.add_paragraph()

    # ========== 一、软件基本信息 ==========
    add_para(doc, "一、软件基本信息", size=12, bold=True)

    table = doc.add_table(rows=0, cols=4)
    table.autofit = False
    widths = [Cm(3.6), Cm(5.2), Cm(3.6), Cm(5.2)]

    def add_row(label1, val1, label2, val2, shade=True):
        row = table.add_row().cells
        for i, w in enumerate(widths):
            row[i].width = w
        fill_cell(row[0], label1, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        fill_cell(row[1], val1)
        fill_cell(row[2], label2, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        fill_cell(row[3], val2)
        if shade:
            shade_cell(row[0])
            shade_cell(row[2])
        for c in row:
            set_cell_border(c, top=True, bottom=True, left=True, right=True)
        return row

    def add_full_row(label, val, shade=True):
        row = table.add_row().cells
        row[0].merge(row[0])
        for i, w in enumerate(widths):
            row[i].width = w
        a = row[0]
        b = row[1].merge(row[2]).merge(row[3])
        fill_cell(a, label, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        fill_cell(b, val)
        if shade:
            shade_cell(a)
        set_cell_border(a, top=True, bottom=True, left=True, right=True)
        set_cell_border(b, top=True, bottom=True, left=True, right=True)
        return a, b

    add_full_row("软件全称", f"{SOFT_FULL}{VERSION}")
    add_full_row("软件简称", SOFT_SHORT)
    add_row("版本号", VERSION, "分类号", "30100（应用软件）")
    add_row("开发完成日期", "2026年08月15日", "首次发表日期", "未发表")
    add_row("发表状态", "未发表", "首次发表地点", "—")
    add_row("开发方式", "单独开发", "软件作品说明", "原创软件")
    add_row("权利取得方式", "原始取得", "权利范围", "全部权利")

    doc.add_paragraph()

    # ========== 二、著作权人信息 ==========
    add_para(doc, "二、著作权人信息", size=12, bold=True)

    t2 = doc.add_table(rows=0, cols=4)
    t2.autofit = False

    def add2_row(label1, val1, label2, val2):
        row = t2.add_row().cells
        for i, w in enumerate(widths):
            row[i].width = w
        fill_cell(row[0], label1, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        fill_cell(row[1], val1)
        fill_cell(row[2], label2, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        fill_cell(row[3], val2)
        shade_cell(row[0])
        shade_cell(row[2])
        for c in row:
            set_cell_border(c, top=True, bottom=True, left=True, right=True)

    def add2_full(label, val):
        row = t2.add_row().cells
        for i, w in enumerate(widths):
            row[i].width = w
        a = row[0]
        b = row[1].merge(row[2]).merge(row[3])
        fill_cell(a, label, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        fill_cell(b, val)
        shade_cell(a)
        set_cell_border(a, top=True, bottom=True, left=True, right=True)
        set_cell_border(b, top=True, bottom=True, left=True, right=True)

    add2_full("著作权人名称", "（请填写申请单位/个人全称，需与证件一致）")
    add2_row("证件类型", "统一社会信用代码证书 / 居民身份证", "证件号码", "（请填写）")
    add2_row("国籍", "中国", "省份/城市", "（请填写省/市）")
    add2_full("通讯地址", "（请填写详细通讯地址及邮政编码）")
    add2_row("联系人", "（请填写）", "联系电话", "（请填写）")
    add2_full("电子邮箱", "（请填写）")

    doc.add_paragraph()

    # ========== 三、软件功能与技术特点 ==========
    add_para(doc, "三、软件功能和技术特点", size=12, bold=True)

    t3 = doc.add_table(rows=0, cols=2)
    t3.autofit = False
    w3 = [Cm(3.6), Cm(14.0)]

    def add3(label, value):
        row = t3.add_row().cells
        row[0].width = w3[0]
        row[1].width = w3[1]
        fill_cell(row[0], label, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        fill_cell(row[1], value)
        shade_cell(row[0])
        set_cell_border(row[0], top=True, bottom=True, left=True, right=True)
        set_cell_border(row[1], top=True, bottom=True, left=True, right=True)

    add3("硬件环境（开发）",
         "Intel Core i7 及以上处理器，16GB 及以上内存，512GB 固态硬盘，100Mbps 及以上网络带宽。")
    add3("硬件环境（运行）",
         "服务端：4 核 CPU、8GB 内存、100GB 硬盘空间；客户端：双核 CPU、4GB 内存、支持 1366×768 及以上分辨率的显示器。")
    add3("软件环境（开发）",
         "操作系统：Windows 10/11 64 位；JDK 17；Node.js 18；Maven 3.9；IntelliJ IDEA 2023、Visual Studio Code；PostgreSQL 14；Redis 7。")
    add3("软件环境（运行）",
         "服务端：CentOS 7.9 / Ubuntu 22.04、JDK 17、PostgreSQL 14、Redis 7、Nginx 1.24、Docker 24；"
         "客户端：Windows 10/11、Edge 110+/Chrome 110+/Firefox 110+；移动端：iOS 14+/Android 10+ 及主流浏览器。")
    add3("编程语言",
         "Java、JavaScript、SQL、HTML、CSS")
    add3("源程序量",
         "65589 行（其中 Java 40330 行、前端 Vue/JavaScript 18176 行、SQL 7083 行）")

    # 主要功能和技术特点
    func_text = (
        "一、开发目的：\n"
        "本软件面向医疗器械及耗材分销行业，针对厂家与经销商在主数据管理、销售采购、仓储库存、"
        "合同授权、手术报台、审批流程、经营分析等环节数据分散、追溯困难、合规压力大等问题，"
        "采用多租户 SaaS 架构，构建覆盖经销商全生命周期的一体化业务管理平台，提升业务协同效率与合规追溯能力。\n"
        "二、主要功能：\n"
        "1. 主数据管理：产品、产品分类、产品线、产品组合、经销商、医院终端、仓库、区域、供应商、产品价格等基础数据维护；\n"
        "2. 合同与授权：合同模板管理、合同在线编制与审批、经销商授权范围与授权期限管理、合同到期提醒；\n"
        "3. 订单业务：销售订单、采购订单、销退订单、采退订单全流程管理，支持订单审批与出库/入库草稿自动生成；\n"
        "4. 仓储库存：收货入库、销售出库、库存移动、库存调整、库存盘点、效期预警、序列号追溯，支持批次与序列号管理；\n"
        "5. 手术与营销：手术植入报台登记、促销规则与起订量/满减策略；\n"
        "6. 审批工作流：可视化审批流配置、顺序/会签/或签节点、加签转办委托、审批监控与超时催办；\n"
        "7. 报表与画像：数据驾驶舱、销售业绩排行、产品销售 TOP10、库存周转、订单追溯、经销商 360 画像；\n"
        "8. 平台后台：多租户开通、租户管理员、角色权限、菜单与按钮配置、数据字典、日志审计；\n"
        "9. 移动端 H5：移动下单、扫码收货、手术报台、移动审批、消息通知；\n"
        "10. 系统管理：用户与角色、数据权限、操作日志、登录日志、接口日志、邮件日志、导入导出。\n"
        "三、技术特点：\n"
        "1. 采用前后端分离架构，后端基于 Spring Boot 3.2 + Java 17 + Spring Security + JPA/MyBatis-Plus，"
        "前端基于 Vue 3 + Vite 5 + Element Plus + Pinia，移动端采用 Vant 4；\n"
        "2. 多租户行级隔离设计，基于 TenantContext 与 Hibernate 拦截器自动注入租户条件，支持厂家租户与经销商租户两级模型；\n"
        "3. 自研审批流引擎，支持模板版本化、条件分支、顺序/会签/或签、前加签后加签、转办委托、抄送与管理员改派终止；\n"
        "4. 库存并发扣减采用乐观条件更新（UPDATE … WHERE qty >= ?）与行级锁，保证批次序列号库存的一致性；\n"
        "5. 业务单号采用 PREFIX-YYYYMMDD-NNNNN 规则，由数据库序列表原子自增生成，杜绝并发重复；\n"
        "6. 全链路操作日志基于 AOP 切面采集，敏感字段自动脱敏，同时写入数据库与按日滚动文件，支持审计导出；\n"
        "7. 报表查询采用 WITH CTE + LEFT JOIN 形式，避免相关子查询导致的数据库异常，提升大数据量下的查询稳定性；\n"
        "8. 基于 Flyway 进行数据库版本化迁移，支持持续集成与多环境一致性部署；\n"
        "9. 基于 Docker Compose 与 Nginx 实现容器化部署，支持测试与生产双环境隔离；\n"
        "10. 接口统一返回 {code, message, data} 结构，业务异常使用错误码枚举，便于前端统一处理与问题定位。"
    )
    add3("主要功能和技术特点", func_text)

    doc.add_paragraph()

    # ========== 四、申请人确认 ==========
    add_para(doc, "四、申请人承诺", size=12, bold=True)
    add_para(doc,
             "申请人保证所提交的申请文件内容真实、合法，申请登记的软件为申请人独立开发并享有完整著作权，"
             "不存在侵犯他人知识产权的情形。如有不实，申请人愿承担由此产生的一切法律责任。",
             size=10.5)
    doc.add_paragraph()
    add_para(doc, "申请人（盖章）：________________________", size=10.5)
    add_para(doc, "经办人（签字）：________________________", size=10.5)
    add_para(doc, "申请日期：       年    月    日", size=10.5)

    doc.save(OUT_FILE)
    print(f"OK -> {OUT_FILE}")


if __name__ == "__main__":
    main()
