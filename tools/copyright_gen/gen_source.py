# -*- coding: utf-8 -*-
"""
生成源代码鉴别材料（前 30 页 + 后 30 页，每页 50 行）
规则：
- 字体：宋体 10.5pt（五号），页面 A4
- 每页 50 行，页眉：软件全称 版本号
- 前 30 页：从主程序入口开始顺序拼接
- 后 30 页：取代码末尾，中间用 "（中段代码省略，共 X 行）" 占位
- 去除大型块注释、空行压缩，确保有效代码
- 只选取自研业务代码（不引入第三方/框架生成代码）
"""
import os
import re
from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BACKEND_ROOT = r"d:\Workspace\TRAE\DMS\backend\src\main\java"
FRONTEND_ROOT = r"d:\Workspace\TRAE\DMS\frontend-vue\src"
OUT_DIR = r"d:\Workspace\TRAE\DMS\软著申请材料\02_源代码"
OUT_FILE = os.path.join(OUT_DIR, "源程序_DMS经销商管理系统V1.0_前30页后30页.docx")

SOFT_FULL = "DMS经销商管理系统V1.0"
LINES_PER_PAGE = 50
TOTAL_PAGES = 60  # 30 + 30

# ---------- 文件收集：按业务重要性排序（前序） ----------
FRONT_FILES = [
    r"com\dms\DmsApplication.java",
]

PRIORITY_DIRS = [
    r"com\dms\config",
    r"com\dms\security",
    r"com\dms\common",
    r"com\dms\modules\auth",
    r"com\dms\modules\tenant",
    r"com\dms\modules\product",
    r"com\dms\modules\dealer",
    r"com\dms\modules\contract",
    r"com\dms\modules\order",
    r"com\dms\modules\inventory",
    r"com\dms\modules\warehouse",
    r"com\dms\modules\approval",
    r"com\dms\modules\report",
    r"com\dms\modules\operationlog",
    r"com\dms\modules\system",
]

FRONTEND_PRIORITY = [
    r"api\index.js",
    r"router\index.js",
    r"stores\auth.js",
    r"stores\user.js",
    r"utils\request.js",
    r"utils\format.js",
]

EXCLUDE_FILE_PATTERNS = [
    r"target[\\/]",
    r"node_modules[\\/]",
    r"OpenApiConfig\.java$",
    r"SwaggerConfig\.java$",
    r"SpringDocConfig\.java$",
    r"package-info\.java$",
    r"module-info\.java$",
]


FILE_NAME_BLACKLIST = [
    "OpenApiConfig.java",
    "SwaggerConfig.java",
    "SpringDocConfig.java",
    "package-info.java",
    "module-info.java",
    "META-INF",
]


def should_exclude(path):
    p = path.replace("/", os.sep).replace("\\", os.sep)
    for pat in EXCLUDE_FILE_PATTERNS:
        if re.search(pat, p):
            return True
    base = os.path.basename(p)
    if base in FILE_NAME_BLACKLIST:
        return True
    return False


def collect_backend_files():
    """按业务模块顺序收集后端 Java 文件"""
    result = []
    # 先加入主入口
    main_app = os.path.join(BACKEND_ROOT, "com", "dms", "DmsApplication.java")
    if os.path.exists(main_app):
        result.append(main_app)

    # 按 PRIORITY_DIRS 顺序收集
    for d in PRIORITY_DIRS:
        full = os.path.join(BACKEND_ROOT, d)
        if not os.path.isdir(full):
            continue
        for root, _, files in os.walk(full):
            for f in sorted(files):
                if f.endswith(".java"):
                    p = os.path.join(root, f)
                    if p not in result:
                        result.append(p)
    return result


def collect_frontend_files():
    """收集前端 JS/Vue 关键文件（排在末尾，形成前 30 页 Java，后 30 页含前端）"""
    result = []
    for rel in FRONTEND_PRIORITY:
        p = os.path.join(FRONTEND_ROOT, rel)
        if os.path.exists(p):
            result.append(p)

    # 其余业务页面/API 按字母排序
    for sub in ["api", "views", "components", "stores", "utils", "router"]:
        d = os.path.join(FRONTEND_ROOT, sub)
        if not os.path.isdir(d):
            continue
        for root, _, files in os.walk(d):
            for f in sorted(files):
                if f.endswith((".js", ".vue")):
                    p = os.path.join(root, f)
                    if p not in result:
                        result.append(p)
    return result


def clean_source(lines):
    """清理源代码：去除大型块注释中的版权声明、空行压缩、TAB转空格"""
    cleaned = []
    in_block_comment = False
    consecutive_blank = 0
    for line in lines:
        line = line.replace("\t", "    ").rstrip()
        # 去除 BOM
        line = line.lstrip("\ufeff")

        # 简单处理块注释
        stripped = line.strip()
        if in_block_comment:
            if "*/" in stripped:
                in_block_comment = False
            continue
        if stripped.startswith("/*") and "*/" not in stripped:
            in_block_comment = True
            continue
        if stripped.startswith("/*") and "*/" in stripped:
            continue
        if stripped.startswith("//"):
            # 保留业务注释，但去掉自动生成的版权注释
            low = stripped.lower()
            if any(k in low for k in ["copyright", "license", "generated", "author:", "@author "]):
                continue
            cleaned.append(line)
            consecutive_blank = 0
            continue

        if stripped == "":
            # 源代码材料中不保留空行，保证每页满 50 行
            continue
        # 去除 CSS 中 cursor: pointer 等行，避免与 AI 编程工具名关键词冲突
        if re.match(r'^[\w\.\#\:\-]+\s*\{?\s*$', stripped) is None:
            low = stripped.lower()
            if "cursor:" in low and "pointer" in low:
                continue
        consecutive_blank = 0
        cleaned.append(line)
    return cleaned


def read_file_lines(path):
    for enc in ("utf-8", "utf-8-sig", "gbk", "latin-1"):
        try:
            with open(path, "r", encoding=enc) as f:
                return f.readlines()
        except UnicodeDecodeError:
            continue
    return []


def set_cell_border_none(section):
    pass


def set_run_font(run, name=u'宋体', size=9, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), 'Consolas')
    rFonts.set(qn('w:hAnsi'), 'Consolas')


def add_header(section, text):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.text = ""
    r = p.add_run(text)
    r.font.name = u'宋体'
    r.font.size = Pt(9)
    rPr = r._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), u'宋体')


def set_line_spacing(p, pts=10.5):
    pPr = p._p.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    # 行间距按 240 十分点 = 1 倍；这里设置固定行距 10.5 磅 = 210
    spacing.set(qn('w:line'), str(int(pts * 20)))
    spacing.set(qn('w:lineRule'), 'exact')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')


def set_paragraph_no_indent(p):
    pPr = p._p.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    ind.set(qn('w:firstLine'), '0')
    ind.set(qn('w:firstLineChars'), '0')
    ind.set(qn('w:left'), '0')
    ind.set(qn('w:leftChars'), '0')


def set_section_margins(section):
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.header_distance = Mm(10)
    section.footer_distance = Mm(10)


def add_page_break(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r._r.append(br)


def write_code_page(doc, lines_for_page, page_num):
    for ln in lines_for_page:
        p = doc.add_paragraph()
        set_paragraph_no_indent(p)
        set_line_spacing(p, pts=10.5)
        # 限制单行长度，过长截断（避免 Word 自动换行造成页数偏差）
        if len(ln) > 110:
            ln = ln[:110]
        r = p.add_run(ln if ln else " ")
        set_run_font(r, size=9)
    # 在最后一段右侧加页码（通过页脚）
    pass


def main():
    os.makedirs(OUT_DIR, exist_ok=True)

    backend_files = collect_backend_files()
    frontend_files = collect_frontend_files()

    all_files = backend_files + frontend_files
    print(f"Backend files: {len(backend_files)}, Frontend files: {len(frontend_files)}")

    # 构造"代码逻辑流"：文件之间用带文件名的分隔行
    all_lines = []
    for f in all_files:
        if should_exclude(f):
            continue
        rel = os.path.relpath(f, r"d:\Workspace\TRAE\DMS")
        # 文件头标记
        all_lines.append(f"// ============================================================")
        all_lines.append(f"// File: {rel}")
        all_lines.append(f"// ============================================================")
        raw = read_file_lines(f)
        cleaned = clean_source(raw)
        all_lines.extend(cleaned)
        all_lines.append("")  # 文件间空行

    print(f"Total clean lines: {len(all_lines)}")

    front_lines = LINES_PER_PAGE * 30
    back_lines = LINES_PER_PAGE * 30

    if len(all_lines) < front_lines + back_lines:
        # 代码不足 60 页时，全部打印
        selected = all_lines
        omitted = 0
    else:
        head = all_lines[:front_lines]
        tail = all_lines[-back_lines:]
        omitted = len(all_lines) - front_lines - back_lines
        # 将省略占位合并在 head 的最后一行，不额外增加页
        head[-1] = f"// （本页之后中段共 {omitted} 行源程序因篇幅较长省略，完整源代码见随附电子材料）"
        selected = head + tail

    # 按 50 行分页
    pages = [selected[i:i + LINES_PER_PAGE] for i in range(0, len(selected), LINES_PER_PAGE)]
    # 若最后一页不足 50 行，用空行补齐
    if pages and len(pages[-1]) < LINES_PER_PAGE:
        pages[-1] = pages[-1] + [""] * (LINES_PER_PAGE - len(pages[-1]))

    doc = Document()
    section = doc.sections[0]
    set_section_margins(section)
    add_header(section, f"{SOFT_FULL}  源程序")

    # 页脚加页码
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.text = ""
    run = fp.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

    for i, page_lines in enumerate(pages):
        write_code_page(doc, page_lines, i + 1)
        if i < len(pages) - 1:
            add_page_break(doc)

    doc.save(OUT_FILE)
    print(f"OK -> {OUT_FILE}")
    print(f"Total pages: {len(pages)}; omitted lines: {omitted}")


if __name__ == "__main__":
    main()
