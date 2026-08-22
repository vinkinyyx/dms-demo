# -*- coding: utf-8 -*-
"""
DMS 通用经销商管理系统 V3.7.1 用户操作手册生成器
- A4 / 宋体正文 / 黑体标题
- 页眉 DMS 通用经销商管理系统 V3.7.1  用户操作手册
- 页脚 第 X 页 共 Y 页
- 内容覆盖 PC 端与移动端全部业务模块，图文一一对应
"""

import os
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from datetime import date

PROJECT_ROOT = r"d:\Workspace\TRAE\DMS"
OUTPUT_DIR = os.path.join(PROJECT_ROOT, "软著申请材料", "03_用户手册")
IMG_DIR = os.path.join(OUTPUT_DIR, "screenshots")
OUTPUT_DOCX = os.path.join(OUTPUT_DIR, "DMS通用经销商管理系统V3.7.1用户操作手册_60页.docx")

HEADER_TEXT = "DMS 通用经销商管理系统 V3.7.1  用户操作手册"
FOOTER_COMPANY = "用户操作手册"

SONG = "宋体"
HEI = "黑体"
KAI = "楷体"


def set_run_font(run, font_name=SONG, size=12, bold=False, color=None):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)


def setup_page(section):
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)


def add_header_footer(section):
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run(HEADER_TEXT)
    set_run_font(hr, SONG, size=9)
    hpPr = hp._p.get_or_add_pPr()
    hpBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '808080')
    hpBdr.append(bottom)
    hpPr.append(hpBdr)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_field(paragraph, instr):
        r = paragraph.add_run()
        fc1 = OxmlElement('w:fldChar'); fc1.set(qn('w:fldCharType'), 'begin')
        it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = instr
        fc2 = OxmlElement('w:fldChar'); fc2.set(qn('w:fldCharType'), 'end')
        r._r.append(fc1); r._r.append(it); r._r.append(fc2)
        return r

    r1 = fp.add_run("第 "); set_run_font(r1, SONG, size=9)
    rf = add_field(fp, "PAGE"); set_run_font(rf, SONG, size=9)
    r2 = fp.add_run(" 页  共 "); set_run_font(r2, SONG, size=9)
    ra = add_field(fp, "NUMPAGES"); set_run_font(ra, SONG, size=9)
    r3 = fp.add_run(" 页"); set_run_font(r3, SONG, size=9)


def set_zh_font(run, name=SONG, size=12, bold=False):
    set_run_font(run, name, size=size, bold=bold)


def add_para(doc, text, size=12, bold=False, font=SONG, align=None, indent=True,
             color=None, space_before=0, space_after=4, line_spacing=1.5):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if indent:
        pf.first_line_indent = Pt(size * 2)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, font, size=size, bold=bold, color=color)
    return p


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    set_run_font(r, HEI, size=18, bold=True)
    p.style = doc.styles['Heading 1']
    for run in p.runs:
        set_run_font(run, HEI, size=18, bold=True)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    p.style = doc.styles['Heading 2']
    r = p.add_run(text)
    set_run_font(r, HEI, size=14, bold=True)
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    p.style = doc.styles['Heading 3']
    r = p.add_run(text)
    set_run_font(r, HEI, size=12, bold=True)
    return p


def add_bullet(doc, text, size=12, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.8 + 0.6 * level)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    set_run_font(r, SONG, size=size)
    return p


def add_number(doc, text, size=12):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    set_run_font(r, SONG, size=size)
    return p


def add_image(doc, img_name, width_cm=14, caption=None):
    img_path = os.path.join(IMG_DIR, img_name)
    if os.path.exists(img_path):
        is_mobile = "_mobile" in img_name or img_name.startswith(("m_", "25_", "26_", "27_", "34_", "35_"))
        w = 7.0 if is_mobile else 15.5
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pPr = p._p.get_or_add_pPr()
        spc = OxmlElement('w:spacing'); spc.set(qn('w:before'), '120'); spc.set(qn('w:after'), '60')
        pPr.append(spc)
        run = p.add_run()
        run.add_picture(img_path, width=Cm(w))
    else:
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell = tbl.rows[0].cells[0]
        cell.width = Cm(width_cm)
        trPr = tbl.rows[0]._tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), '3600')
        trHeight.set(qn('w:hRule'), 'atLeast')
        trPr.append(trHeight)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"［ 系统截图位置：{img_name} ］")
        set_run_font(r, size=10, color=RGBColor(0x80, 0x80, 0x80))
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(caption)
        set_run_font(cr, size=9)


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), '18'); left.set(qn('w:space'), '8'); left.set(qn('w:color'), '4472C4')
    pBdr.append(left)
    pPr.append(pBdr)
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'F2F6FC')
    pPr.append(shd)
    r = p.add_run("【说明】" + text)
    set_run_font(r, SONG, size=10.5, color=RGBColor(0x33, 0x33, 0x33))


def style_table(table):
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    set_run_font(run, SONG, size=10.5)


def add_table_kv(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, HEI, size=10.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), '4472C4')
        tcPr.append(shd)
    for i, row in enumerate(rows, start=1):
        cells = table.rows[i].cells
        for j, val in enumerate(row):
            cells[j].text = ""
            p = cells[j].paragraphs[0]
            r = p.add_run(str(val))
            set_run_font(r, SONG, size=10.5)
            cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.style = 'Table Grid'
    if col_widths:
        for row in table.rows:
            for j, w in enumerate(col_widths):
                row.cells[j].width = Cm(w)
    return table


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


# =============================================================================
def build_cover(doc):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DMS 通用经销商管理系统"); set_run_font(r, HEI, size=36, bold=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("V3.7.1"); set_run_font(r, HEI, size=24, bold=True, color=RGBColor(0x44, 0x72, 0xC4))
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("用 户 操 作 手 册"); set_run_font(r, HEI, size=30, bold=True)
    for _ in range(6):
        doc.add_paragraph()
    info = [
        ("文档版本", "V3.7.1"),
        ("适用平台", "PC 端（Vue 3 + Element Plus）、移动端 H5（Vant 4）"),
        ("适用读者", "系统管理员、业务操作员、经销商、仓库管理员、质检员、财务人员"),
        ("编制日期", date.today().strftime("%Y 年 %m 月 %d 日")),
    ]
    t = doc.add_table(rows=len(info), cols=2); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(info):
        c0 = t.rows[i].cells[0]; c1 = t.rows[i].cells[1]
        c0.text = ""; c1.text = ""
        r0 = c0.paragraphs[0].add_run(k); set_run_font(r0, HEI, size=12, bold=True)
        c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = c1.paragraphs[0].add_run(v); set_run_font(r1, SONG, size=12)
        c0.width = Cm(3.5); c1.width = Cm(10)
    page_break(doc)


def build_toc(doc):
    add_h1(doc, "目  录")
    items = [
        ("第一章  手册概述", "1"),
        ("第二章  系统运行环境与登录", "2"),
        ("第三章  工作台与系统首页", "4"),
        ("第四章  基础数据管理", "6"),
        ("第五章  合同与授权管理", "12"),
        ("第六章  销售订单与销售出库", "15"),
        ("第七章  采购订单与采购入库", "19"),
        ("第八章  库存业务管理", "22"),
        ("第九章  产品追溯与序列号管理", "26"),
        ("第十章  手术报告与市场营销", "28"),
        ("第十一章  审批中心与工作流", "31"),
        ("第十二章  数据看板与业务报表", "34"),
        ("第十三章  移动端 H5 操作说明", "38"),
        ("第十四章  系统管理与配置", "44"),
        ("第十五章  消息中心与日志审计", "48"),
        ("第十六章  常见问题与故障排查", "50"),
        ("第十七章  附录：快捷键、术语表与版本记录", "52"),
    ]
    for title, page in items:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.6
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(14.5), WD_ALIGN_PARAGRAPH.RIGHT, 2)
        r = p.add_run(title); set_run_font(r, SONG, size=12)
        r2 = p.add_run("\t" + page); set_run_font(r2, SONG, size=12)
    page_break(doc)


# ----------- 第一章 -----------
def chapter1(doc):
    add_h1(doc, "第一章  手册概述")
    add_h2(doc, "1.1  编写目的")
    add_para(doc, "本手册面向 DMS 通用经销商管理系统 V3.7.1（以下简称“本系统”或“DMS”）的所有使用者，"
                  "用于指导业务人员完成日常业务操作，指导系统管理员完成系统配置与维护，"
                  "同时也可供企业培训、内部审计及软件著作权登记参考。")
    add_para(doc, "本手册按照“先 PC、后移动”、“先基础后业务”、“先操作后管理”的顺序编排，"
                  "每一个操作步骤均配有系统真实界面截图，图注与图中显示的功能名称、按钮、字段完全一致，"
                  "便于读者按图索骥、对照操作。")
    add_h2(doc, "1.2  适用对象")
    add_bullet(doc, "系统管理员：负责租户、账号、角色、菜单、数据字典及业务参数配置。")
    add_bullet(doc, "销售业务员：负责经销商档案、销售订单、合同与授权的录入与跟踪。")
    add_bullet(doc, "采购业务员：负责采购订单、采购收货、采购退货等业务。")
    add_bullet(doc, "仓库管理员：负责出入库、库存移动、库存盘点、库存调拨、效期预警、扫码作业。")
    add_bullet(doc, "质量管理员：负责产品资质、UDI 追溯、序列号追溯、批次追溯、合规审计。")
    add_bullet(doc, "财务人员：负责应收账款、收款核销、销售/采购金额统计、报表订阅。")
    add_bullet(doc, "经销商/医院终端用户：通过移动端 H5 进行下单、查单、手术报台、查看授权。")
    add_h2(doc, "1.3  功能概览")
    add_para(doc, "DMS 通用经销商管理系统覆盖医疗器械经销业务的全链条，主要功能模块如下表所示：")
    add_table_kv(doc,
        ["业务域", "核心模块", "主要功能"],
        [
            ["基础数据", "产品、分类、仓库、经销商、医院、供应商", "档案维护、资质管理、价格策略、GSP/GMP 证照预警"],
            ["合同授权", "合同工作台、合同模板、授权管理", "合同审批、电子归档、授权经销商与产品范围控制"],
            ["销售业务", "销售订单、销售出库、销退订单", "订单审批、扫码出库、发票、收货、销退冲红"],
            ["采购业务", "采购订单、采购入库、采购退货", "采购审批、收货质检、批次/序列号录入"],
            ["库存业务", "库存查询、库存移动、库存调整、库存盘点、效期预警", "多仓多货位、批次/序列号、先进先出、近效期预警"],
            ["追溯合规", "序列号追溯、批次追溯、UDI 追溯、操作日志", "一物一码、来源去向、操作留痕、合规报表"],
            ["手术营销", "手术报告、促销活动、费用核销", "植入报台、跟台记录、促销政策、费用闭环"],
            ["审批中心", "审批模板、审批任务、审批监控", "可视化流程、多级审批、超时提醒、代理审批"],
            ["数据报表", "工作台、数据看板、销售/库存/采购报表", "实时图表、TOP 排行、周转率、订单追溯、导出订阅"],
            ["移动作业", "移动下单、移动审批、扫码收货、扫码盘点、手术报台", "PDA/手机 H5 全覆盖，离线缓存"],
            ["系统管理", "账号、角色、菜单、数据字典、列表页配置、参数", "租户隔离、细粒度权限、审计日志、可配置界面"],
        ],
        col_widths=[2.5, 4.5, 8.0])
    add_h2(doc, "1.4  术语约定")
    add_table_kv(doc,
        ["术语", "含义"],
        [
            ["租户", "一套独立的业务数据空间，数据与其他租户物理或逻辑隔离，本系统通过 tenantCode 区分。"],
            ["SKU", "最小库存单元，对应一个具体的产品规格型号。"],
            ["UDI", "医疗器械唯一标识（Unique Device Identification），由 DI + PI 组成。"],
            ["批号", "同一批次生产或采购的产品标识，用于批次追溯与效期管理。"],
            ["序列号", "单件产品唯一编号，用于高值耗材/植入物的一物一码管理。"],
            ["GSP/GMP", "药品/医疗器械经营/生产质量管理规范，系统对相关证照进行效期预警。"],
            ["工作流", "审批模板定义的多级审批节点、条件分支与通知规则。"],
        ],
        col_widths=[3.0, 12.0])
    page_break(doc)


# ----------- 第二章 -----------
def chapter2(doc):
    add_h1(doc, "第二章  系统运行环境与登录")
    add_h2(doc, "2.1  运行环境")
    add_h3(doc, "2.1.1  服务端环境")
    add_bullet(doc, "操作系统：Linux（推荐 CentOS 7+ / Ubuntu 20.04+），运行 Docker 24+ 与 Docker Compose。")
    add_bullet(doc, "JDK：Eclipse Temurin 17；应用服务器：Spring Boot 3.2 内嵌 Tomcat。")
    add_bullet(doc, "数据库：PostgreSQL 14+；缓存：Redis 7；消息：异步线程池 + 数据库事件表。")
    add_bullet(doc, "反向代理：Nginx 1.24，提供 PC 与移动端 H5 静态资源及 API 代理。")
    add_h3(doc, "2.1.2  客户端环境")
    add_table_kv(doc,
        ["端", "推荐浏览器 / 设备", "最低分辨率"],
        [
            ["PC 端", "Chrome 100+、Edge 100+、Firefox 100+（均为近一年版本）", "1366 × 768，推荐 1920 × 1080"],
            ["移动端 H5", "iOS Safari 15+、Android Chrome 100+、企业微信/钉钉/飞书内置浏览器", "375 × 667，推荐 390 × 844"],
            ["扫码设备", "支持摄像头的手机/PDA，或接入 USB 扫码枪的 PC", "—"],
        ],
        col_widths=[2.5, 9.0, 4.0])
    add_h2(doc, "2.2  PC 端登录")
    add_para(doc, "打开浏览器，在地址栏输入系统访问地址（正式环境示例：http://8.133.193.238:8081/，"
                  "测试环境示例：http://8.133.193.238:8083/），进入系统登录页。")
    add_image(doc, "01_login.png", caption="图 2-1  PC 端登录页")
    add_para(doc, "登录页包含以下输入项：")
    add_table_kv(doc,
        ["字段", "是否必填", "说明"],
        [
            ["租户", "是", "由系统管理员分配的租户编码，如 default。"],
            ["账号", "是", "用户登录账号，如 admin 或 sys_admin。"],
            ["密码", "是", "登录密码，密码长度不少于 8 位，需包含字母与数字。"],
            ["验证码", "视策略", "如管理员开启了验证码策略，则必须填写图形验证码。"],
            ["记住我", "否", "勾选后 7 天内免重新登录（公用电脑请勿勾选）。"],
        ],
        col_widths=[2.5, 2.0, 10.5])
    add_para(doc, "登录操作步骤：")
    add_number(doc, "在“租户”输入框输入租户编码，按 Tab 键或用鼠标点击下一输入框。")
    add_number(doc, "依次输入“账号”“密码”，必要时输入验证码。")
    add_number(doc, "单击蓝色“登 录”按钮，系统校验通过后跳转至工作台首页；校验失败时在按钮下方显示红色提示。")
    add_note(doc, "连续 5 次输入错误密码，账号将被临时锁定 15 分钟；忘记密码请联系系统管理员在“账号管理”中重置。")
    add_h2(doc, "2.3  移动端登录")
    add_para(doc, "在手机浏览器中输入移动端地址（如 http://8.133.193.238:8081/mobile/login），"
                  "或通过企业微信/钉钉工作台进入，进入移动端登录页。移动端登录页采用卡片式布局，"
                  "自上而下依次为租户、账号、密码三个输入框及蓝色“登录”按钮。")
    add_image(doc, "25_mobile_login.png", caption="图 2-2  移动端登录页")
    add_para(doc, "移动端登录成功后，系统默认进入“首页”，底部 Tab 栏提供“首页、订单、报台、审批、消息、我的”六个一级入口。"
                  "Token 有效期默认 2 小时，刷新 Token 有效期 7 天，过期后自动跳回登录页。")
    add_h2(doc, "2.4  密码修改与退出登录")
    add_para(doc, "PC 端：右上角头像下拉菜单中选择“修改密码”，输入旧密码与新密码后保存；选择“退出登录”可清除本机会话。")
    add_para(doc, "移动端：在底部 Tab 栏切换到“我的”，点击“修改密码”或“退出登录”即可完成相应操作。")
    page_break(doc)


# ----------- 第三章 -----------
def chapter3(doc):
    add_h1(doc, "第三章  工作台与系统首页")
    add_h2(doc, "3.1  PC 端工作台首页")
    add_para(doc, "PC 端登录成功后默认进入“工作台首页”。页面顶部横向显示“欢迎使用 DMS 通用经销商管理系统”标题与"
                  "“当前登录：(用户)，租户 -”副标题，右上角有“Vue3 + Element Plus 栈”技术标识与“仪表盘数据：本年”周期标签。"
                  "下方依次为四个快捷入口卡片（产品管理、经销商管理、销售订单、库存查询）、仪表盘速览 KPI、"
                  "销售趋势折线图、订单漏斗图、经销商销售 TOP5 与本月速览。")
    add_image(doc, "38_init_wizard.png", caption="图 3-1  PC 端工作台首页")
    add_h3(doc, "3.1.1  欢迎区与周期切换")
    add_para(doc, "欢迎区位于首页顶部，左侧为 DMS 蓝色 logo + “欢迎使用 DMS 通用经销商管理系统”标题，"
                  "右侧显示当前登录用户与租户；“仪表盘速览”卡片右上角提供四个时间段按钮：当日、本月、本季、本年，"
                  "默认选中“本年”（蓝色高亮），点击任一按钮所有 KPI 与图表即时切换；旁边还有“刷新”按钮与"
                  "“查看完整仪表盘 →”链接，可跳转到数据看板。")
    add_h3(doc, "3.1.2  快捷入口")
    add_para(doc, "欢迎区下方为四个常用功能卡片，自左向右依次为：产品管理（蓝色购物袋图标）、"
                  "经销商管理（蓝色楼宇图标）、销售订单（蓝色购物车+出箭头图标）、库存查询（蓝色仓库图标）。"
                  "单击任一卡片即可直达对应模块，减少菜单层级，鼠标悬停时卡片有轻微上浮阴影效果。")
    add_h3(doc, "3.1.3  仪表盘速览 KPI")
    add_para(doc, "仪表盘速览区包含四张 KPI 卡片，顶部为彩色横条，数值与说明文字居中：")
    add_table_kv(doc,
        ["指标", "示例值", "颜色"],
        [
            ["销售总额", "¥ 8,705,187.46", "蓝色"],
            ["订单数", "695", "绿色"],
            ["活跃经销商", "50", "橙色"],
            ["手术台数", "30", "红色"],
        ],
        col_widths=[4.0, 6.0, 5.0])
    add_h3(doc, "3.1.4  销售趋势与订单漏斗")
    add_para(doc, "“销售趋势（本年）”为蓝色面积折线图，横轴为月份 2026-01 至 2026-08，纵轴为金额（0–300,000 等刻度），"
                  "曲线展示年度内销售金额走势，悬停可查看当月具体数值。"
                  "右侧“订单漏斗”自上而下按状态分层：APPROVED（已审批，黄色最宽）、SUBMITTED（已提交，绿色）、"
                  "DRAFT（草稿，深蓝）、COMPLETED（已完成，红色）、CANCELLED（已取消，浅蓝）、"
                  "REJECTED（已驳回，深绿）、SHIPPING（发货中，橙色），直观反映订单各阶段留存。")
    add_para(doc, "页面再下方依次展示“经销商销售 TOP5”排行榜与“本月速览”待办/统计卡，"
                  "为管理人员提供一屏式业务总览。")
    add_h2(doc, "3.2  仪表盘")
    add_para(doc, "在左侧菜单进入“数据看板 → 仪表盘”，可查看完整的业务仪表盘。该页面在工作台卡片基础上扩展了"
                  "产品销售 TOP10、经销商销售排行、库存周转率趋势、订单审批时效等图表，并支持按时间区间、经销商、产品分类筛选。")
    add_image(doc, "21_bi.png", caption="图 3-2  数据看板仪表盘")
    add_para(doc, "图表操作要点：")
    add_bullet(doc, "鼠标悬停在图表数据点上可查看具体数值；点击图例可隐藏/显示对应系列。")
    add_bullet(doc, "图表右上角提供“下载为图片”“导出 Excel”按钮，便于线下汇报。")
    add_bullet(doc, "筛选条件变更后图表自动刷新，无需手动点击查询。")
    page_break(doc)


# ----------- 第四章 -----------
def chapter4(doc):
    add_h1(doc, "第四章  基础数据管理")
    add_para(doc, "基础数据是 DMS 所有业务单据的前置条件，必须在开展销售、采购、库存业务前完成维护。"
                  "本章按菜单顺序依次说明产品管理、产品分类、仓库管理、经销商管理、医院/终端、供应商等模块。")

    add_h2(doc, "4.1  产品管理")
    add_h3(doc, "4.1.1  产品列表")
    add_para(doc, "进入“基础数据 → 产品管理”，打开产品列表页。页面顶部提供关键词、产品分类、产品类型、状态四个筛选条件，"
                  "工具栏提供“新增、导入、导出、批量删除”按钮，表格列依次为产品编码、中文名称、英文名称、产品类型、规格型号、"
                  "单位、参考单价、状态等。")
    add_image(doc, "04_product_list.png", caption="图 4-1  产品列表")
    add_para(doc, "列表常用操作：")
    add_bullet(doc, "查询：输入关键词（产品编码/中文名/英文名/拼音码）后点击“查询”，表格按条件刷新。")
    add_bullet(doc, "重置：清空所有筛选条件并刷新表格。")
    add_bullet(doc, "新增：点击右上角“新增”按钮弹出新增产品表单。")
    add_bullet(doc, "查看/编辑：操作列点击“查看”进入只读详情，点击“更多 → 编辑”进入编辑表单。")
    add_bullet(doc, "导入：下载模板，填写后批量上传，系统逐行校验并返回成功/失败明细。")
    add_bullet(doc, "导出：按当前筛选条件导出 Excel，包含所有列。")

    add_h3(doc, "4.1.2  新增/编辑产品")
    add_image(doc, "05_product_edit.png", caption="图 4-2  新增产品表单")
    add_para(doc, "新增产品表单字段说明如下：")
    add_table_kv(doc,
        ["字段", "必填", "说明"],
        [
            ["产品编码", "是", "企业内唯一编码，建议规则：分类码(2) + 序号(6)，保存后不可修改。"],
            ["中文名称", "是", "产品注册证上的中文通用名称。"],
            ["英文名称", "否", "产品英文名称，用于出口单据。"],
            ["产品类型", "是", "下拉选择：普通耗材 / 高值耗材 / 植入物 / 设备 / 试剂。"],
            ["产品分类", "是", "树形选择，来源于“产品分类”模块。"],
            ["规格型号", "是", "具体规格与型号。"],
            ["单位", "是", "计量单位，如 个、盒、支、套。"],
            ["参考单价", "否", "默认销售单价，可在订单中按经销商价格策略覆盖。"],
            ["税率", "是", "下拉选择 0%、3%、6%、9%、13% 等。"],
            ["UDI 追溯", "否", "开关，开启后出入库必须采集 UDI 编码。"],
            ["序列号管理", "否", "开关，开启后必须按单件序列号管理库存。"],
            ["临期预警月", "否", "距离失效日期剩余多少个月时触发近效期预警。"],
            ["安全库存", "否", "低于该数量时在库存查询中高亮提醒。"],
            ["最小订购量", "否", "销售/采购下单的最小数量，低于此值系统拒绝提交。"],
            ["状态", "是", "启用 / 停用，停用产品不能在新单据中选择。"],
        ],
        col_widths=[3.0, 1.5, 10.5])
    add_para(doc, "保存后系统自动校验编码唯一性，若编码重复则在字段下方提示“产品编码已存在”。")

    add_h2(doc, "4.2  产品分类")
    add_para(doc, "进入“基础数据 → 产品分类”，以平铺列表方式维护产品分类档案。"
                  "顶部筛选区提供关键词输入框及“查询、重置”按钮，工具栏右侧为“批量删除、新增”，"
                  "表格列依次为：编号、分类编码（如 CAT-DEV 器械、CAT-CON 介入耗材、CAT-REA 骨科植入、"
                  "CAT-MON 监护设备、CAT-REH 康复器械及其下级编码 CAT-DEV-01 心血管、CAT-CON-01 手术室耗材等）、"
                  "分类名称、排序（数值越小越靠前）、状态（绿色“启用”标签）、创建时间、更新时间与操作列，"
                  "操作列提供“查看”按钮和“更多”下拉（编辑、删除等）。底部分页区显示总条数（如共 10 条）、当前页与每页条数。")
    add_image(doc, "06_category.png", caption="图 4-3  产品分类")
    add_para(doc, "新增分类时需录入分类编码、分类名称、上级分类（可留空表示根分类）、排序值与状态，"
                  "其中编码全局唯一；若所选上级分类被停用，则子分类在新单据中同样不可选。"
                  "删除分类前系统会校验该分类下是否仍关联产品，若存在关联则拒绝删除并提示先迁移产品。")

    add_h2(doc, "4.3  仓库管理")
    add_para(doc, "进入“基础数据 → 仓库管理”，维护企业的物理仓、医院寄售仓、虚拟仓等。"
                  "顶部筛选区为关键词、状态，工具栏为“查询、重置、批量删除、新增”。"
                  "表格列依次为：编号、仓库编码（WH-MAIN 主仓 / WH-SUB 子仓 / WH-HP 医院仓）、仓库名称（如“一级经销商-1-主仓”）、"
                  "类型（main 主仓 / sub 子仓 / hospital 医院仓）、状态（启用/停用）、创建时间、更新时间、操作。")
    add_image(doc, "08_warehouse.png", caption="图 4-4  仓库管理")
    add_para(doc, "点击“新增”创建仓库时需录入仓库编码、仓库名称、仓库类型、负责人、联系电话、地址、是否启用库位管理等信息。"
                  "启用库位管理后，库存必须精确到货位编号；不启用则按仓库汇总数量。"
                  "操作列“查看”按钮以只读方式打开仓库详情，“更多”下拉中提供编辑、停用/启用等操作。")

    add_h2(doc, "4.4  经销商管理")
    add_h3(doc, "4.4.1  经销商档案")
    add_para(doc, "进入“基础数据 → 经销商管理”，维护经销商档案。顶部筛选区提供关键词、经销商级别（一级/二级）、"
                  "GSP 资质状态（active/inactive）、状态等条件；工具栏提供“查询、重置、批量删除、新增”按钮。"
                  "表格列依次为：编号、经销商编码（D00001 等）、经销商名称、经销商级别、GSP 状态、GSP 有效期、"
                  "法人、联系人、联系电话、邮箱、状态、创建时间、更新时间、操作。")
    add_image(doc, "39_org.png", caption="图 4-5  经销商档案列表")
    add_para(doc, "经销商档案字段包括：基础信息（编码、名称、简称、级别、上级经销商）、证照信息（营业执照、"
                  "医疗器械经营许可证、GSP 认证，均支持上传附件与到期日）、联系人、收货地址、开票信息、银行账户、"
                  "价格策略、信用额度、账期等。操作列“查看”按钮以只读方式打开经销商详情，"
                  "“更多”下拉中提供编辑、停用/启用、重置密码等操作。")
    add_h3(doc, "4.4.2  经销商画像列表")
    add_para(doc, "进入“数据看板 → 经销商画像”，以列表形式概览全部经销商的关键经营指标。"
                  "表格列依次为：经销商编码、经销商名称、级别、本月销售额（橙色高亮）、YTD 销售额（年度累计）、"
                  "同比增长率、库存数量、应收账款、最近下单日期、操作。"
                  "操作列的“查看画像”按钮可进入该经销商的 360 画像详情页。")
    add_image(doc, "07_dealer.png", caption="图 4-6  经销商画像列表")
    add_h3(doc, "4.4.3  经销商 360 画像")
    add_para(doc, "在经销商画像列表点击“查看画像”按钮，进入该经销商的 360 画像详情页。"
                  "页面左侧为经销商基础信息卡（编码、名称、等级、法人、税号、GSP 状态、GSP 有效期、联系人、联系电话、邮箱）；"
                  "右侧顶部为 KPI 概览区，包含三张主指标卡（本月销售目标达成、年累计目标达成、环比上月）"
                  "与六张副指标卡（本月订单数、YTD 净返利、YTD 退货、库存数量、库存 SKU、有效合同数）。")
    add_image(doc, "22_dealer_profile.png", caption="图 4-7  经销商 360 画像")
    add_para(doc, "KPI 区下方提供五个标签页：KPI 概览、月度达成、返利明细、合同列表、库存明细。"
                  "“月度达成”标签页展示“月度目标达成趋势”组合图——蓝色柱表示实际销售额、虚线表示目标、"
                  "绿色折线表示达成率（右轴），横轴为 2025-01 至 2026-12 的月份。"
                  "页面左上角“← 返回”按钮可返回画像列表。")
    add_note(doc, "GSP 资质即将到期（默认 30 天内）时，系统在工作台“待办事项”与消息中心同步提醒。")

    add_h2(doc, "4.5  医院/终端与供应商")
    add_para(doc, "医院/终端模块维护医疗机构档案，字段与经销商类似，重点维护医院等级、科室、跟台销售员等；"
                  "供应商模块维护供货厂家、一级代理等上游单位，字段包括生产许可证、医疗器械注册证及其有效期。"
                  "两类档案的新增、编辑、导入、导出操作与经销商完全一致。")

    add_h2(doc, "4.6  账号管理")
    add_para(doc, "进入“用户与权限 → 账号管理”，维护系统用户账号。"
                  "顶部提供关键词、状态两个筛选条件及“查询、重置”按钮，右上角为“新增”按钮。"
                  "数据表格列依次为编号、账号（如 sys_admin、admin、vendor04、dealer01 等）、姓名、角色（系统管理员/销售/经销商管理员）、"
                  "类型（厂商 / 经销商）、邮箱、手机号、状态（绿色“启用”标签，另有停用、锁定两种状态）、创建时间、更新时间与操作列，"
                  "操作列提供“查看”按钮和“更多”下拉。底部分页区显示总条数与每页条数。")
    add_image(doc, "03_profile.png", caption="图 4-8  账号管理")
    add_para(doc, "新增账号时需录入账号、姓名、用户类型（厂商 / 经销商）、初始密码（至少 8 位）、角色、邮箱、手机号、"
                  "所属组织、所属经销商（经销商类型账号）与状态。保存后系统可通过邮件或短信发送初始密码。"
                  "管理员可对账号执行启用 / 停用 / 锁定、重置密码、删除等操作。")
    page_break(doc)


# ----------- 第五章 -----------
def chapter5(doc):
    add_h1(doc, "第五章  合同与授权管理")
    add_h2(doc, "5.1  合同工作台")
    add_para(doc, "进入“合同管理 → 合同工作台”，按状态标签页（全部、草稿、审批中、已生效、已驳回、已终止、已到期）"
                  "分类查看合同。顶部支持按合同编号/名称、合同分类筛选，工具栏提供“新建合同、导出、异步导出”按钮。")
    add_image(doc, "09_contract.png", caption="图 5-1  合同工作台")
    add_para(doc, "新建合同操作步骤：")
    add_number(doc, "点击“新建合同”按钮，在弹出的表单中选择合同分类（年度框架协议、单笔销售合同、寄售协议等）。")
    add_number(doc, "填写合同编号（可自动生成）、合同名称、经销商、签订日期、生效日期、到期日期、合同金额。")
    add_number(doc, "在“合同产品”子表中添加授权产品清单，逐行选择产品并约定单价、数量、返点政策。")
    add_number(doc, "上传合同扫描件（PDF/图片），提交审批。")
    add_number(doc, "审批通过后合同状态变为“已生效”，可在销售订单中被引用；到期前 30 天系统自动提醒续签。")
    add_h2(doc, "5.2  合同模板")
    add_para(doc, "进入“合同管理 → 合同模板”，管理员可维护标准合同模板，设定占位符（如 ${dealerName}、${amount}），"
                  "新建合同时基于模板自动生成合同正文，减少重复录入。")
    add_h2(doc, "5.3  授权管理")
    add_para(doc, "进入“合同管理 → 授权管理”，维护厂商授予经销商的销售授权书。"
                  "顶部提供“关键词”搜索框与“状态”下拉筛选，右侧蓝色“新增”按钮用于新建授权。"
                  "表格列依次为：编号、授权编号、经销商、生效日期、截止日期、状态（启用/停用）、创建时间、更新时间、操作。")
    add_image(doc, "10_authorization.png", caption="图 5-2  经销商授权管理")
    add_para(doc, "操作列“查看”按钮以只读方式打开授权详情，“更多”下拉中提供编辑、续签、作废等操作。"
                  "新增授权时需选择经销商、授权产品范围（按产品分类或具体产品勾选）、授权区域、授权起止日期并上传授权书扫描件。"
                  "保存后状态默认为“启用”，到期后自动失效，销售下单时系统自动校验授权范围与有效期。")
    page_break(doc)


# ----------- 第六章 -----------
def chapter6(doc):
    add_h1(doc, "第六章  销售订单与销售出库")
    add_h2(doc, "6.1  销售订单列表")
    add_para(doc, "进入“订单业务 → 销售订单”，打开销售订单列表。顶部筛选区依次为关键词、状态、开始日期、结束日期、经销商，"
                  "工具栏按钮依次为“查询、重置、导入、导出、新增”。"
                  "表格列依次为：编号、销售单号（SO-xxxxxxxxxx）、类型（PURCHASE 采购补货 / EMERGENCY 紧急 / SHORTAGE 缺货 / CUSTOM 定制等）、"
                  "经销商、发货仓库、含税金额、实付金额、状态（草稿/已提交/已审批等彩色标签）、审核人、创建时间、操作。")
    add_image(doc, "11_sales_order.png", caption="图 6-1  销售订单列表")
    add_para(doc, "操作列“查看”按钮打开订单详情弹窗，“更多”下拉中提供编辑、提交审批、取消、打印等操作。"
                  "右上角“新增”按钮进入下单页；“导入”支持按模板批量导入订单，“导出”按当前筛选条件导出 Excel。")
    add_h2(doc, "6.2  订单详情")
    add_para(doc, "在订单列表点击操作列“详情”按钮，打开订单详情弹窗。弹窗自上而下分为三个区域：")
    add_image(doc, "11_order_detail.png", caption="图 6-2  销售订单详情")
    add_table_kv(doc,
        ["区域", "显示内容"],
        [
            ["详情（基础信息）", "以两列栅格展示：编号、编码（SO-xxxxxxxxxx）、订单类型、经销商、含税金额、优惠金额、最终金额、税额、期望到货、状态（草稿/已提交/已审批等）、备注、创建时间、提交时间。"],
            ["明细（订单行）", "表格列依次为：顺序、产品、productCode（产品编码）、产品名称、productSpec（规格型号）、数量、unitPrice（单价）、税率、subtotal（小计）、isGift（是否赠品）。"],
            ["操作记录", "表格列：操作人、操作、变更内容、操作时间，按时间倒序记录订单每次状态变更与字段修改；若暂无操作记录则显示“暂无数据/暂无操作记录”占位。"],
        ],
        col_widths=[3.0, 12.0])
    add_para(doc, "点击弹窗右上角“×”关闭详情返回列表。订单的收货信息、发票信息、审批记录与出库记录可在订单编辑页或后续出库单中查看。")
    add_h2(doc, "6.3  销售下单操作")
    add_para(doc, "在销售订单列表点击“新增订单”，按以下步骤操作：")
    add_number(doc, "选择经销商。系统自动带出该经销商的合同、授权、价格策略及收货地址。")
    add_number(doc, "在订单行中点击“添加产品”，通过编码/名称/扫码选择产品，输入数量与单价。")
    add_number(doc, "系统按税率自动计算金额、税额、价税合计；若超出经销商信用额度，提示但允许提交（由审批节点决定）。")
    add_number(doc, "维护收货地址、发票信息、备注与附件。")
    add_number(doc, "点击“保存草稿”或“提交审批”。提交后按审批模板流转，审批通过后进入“已审批”状态，等待出库。")
    add_note(doc, "若所选产品不在经销商授权范围内，或授权已过期，系统阻止提交并提示具体原因。")
    add_h2(doc, "6.4  销售出库")
    add_para(doc, "进入“库存业务 → 销售出库”，对已审批订单执行出库作业。页面顶部筛选区依次为关键词、状态、开始日期、结束日期，"
                  "工具栏按钮为“查询、重置、导入、导出”。表格列依次为：编号、出库单号（GI-YYYYMMDD-xxxxx）、经销商、仓库、"
                  "来源订单（SO- 单据号）、状态（草稿/部分发货等彩色标签）、创建时间、更新时间、操作。"
                  "操作列提供“查看”与“更多”两个按钮。")
    add_image(doc, "12_sales_out.png", caption="图 6-3  销售出库单")
    add_para(doc, "出库操作步骤：")
    add_number(doc, "点击“新增出库”，选择来源销售订单，系统自动带出待出库产品与数量。")
    add_number(doc, "选择发货仓库，逐行录入批号/序列号（启用了批次/序列号管理的产品必须录入）。")
    add_number(doc, "支持扫码：将光标定位到“扫码输入框”，用扫码枪扫描产品条码或 UDI，系统自动累加数量并校验库存。")
    add_number(doc, "核对无误后点击“确认出库”，系统扣减库存并生成库存事务记录；部分发货时来源订单自动保留剩余数量，可继续出库。")
    add_number(doc, "打印随货同行单与出库复核单，交接给物流。")
    add_h2(doc, "6.5  销售退货")
    add_para(doc, "进入“订单业务 → 销退订单”，对已出库订单进行退货处理。"
                  "新建销退单时选择原销售出库单，逐行录入退货数量与原因，保存后提交审批；"
                  "审批通过后仓库执行收货入库，库存增加并冲减原销售金额。")
    add_image(doc, "28_sales_return.png", caption="图 6-4  销退订单")
    page_break(doc)


# ----------- 第七章 -----------
def chapter7(doc):
    add_h1(doc, "第七章  采购订单与采购入库")
    add_h2(doc, "7.1  采购订单")
    add_para(doc, "进入“订单业务 → 采购订单”，维护向供应商下达的采购订单。页面顶部筛选区依次为关键词、状态、开始日期、结束日期、供应商，"
                  "工具栏按钮依次为“查询、重置、导入、导出、新增”。"
                  "表格列依次为：编号、采购单号（PO- 前缀自动生成）、类型（普通采购/紧急补货等）、供应商、入库仓库、"
                  "总金额、实付金额、状态、审核人、创建时间、操作。当前筛选条件下若无记录，表格区域显示“暂无数据”占位，"
                  "底部分页区显示“共 0 条”。")
    add_image(doc, "13_purchase.png", caption="图 7-1  采购订单列表（暂无数据状态）")
    add_para(doc, "新建采购订单时需选择供应商、收货仓库，添加采购产品行（产品、数量、单价、税率、预计到货日期），"
                  "可维护付款方式、交货条款、备注附件。保存后订单状态默认为“草稿”，提交后按采购审批流程流转，"
                  "审批通过后可在“收货入库”中被引用。")
    add_h2(doc, "7.2  采购收货")
    add_para(doc, "进入“库存业务 → 收货入库”，对已审批采购订单执行收货。"
                  "选择采购订单后系统带出待收明细，仓库收货时录入实际到货数量、批号、生产日期、失效日期、序列号；"
                  "若启用了质检流程，收货后状态为“待检”，质检员在“库存调整/质检”中放行或拒收。")
    add_image(doc, "14_receipt.png", caption="图 7-2  采购收货单")
    add_h2(doc, "7.3  采购报表")
    add_para(doc, "在“业务报表 → 报表中心”中可查看采购相关报表（如采购明细、采购统计），按供应商、产品、日期范围统计采购金额、"
                  "采购数量、到货及时率等指标，支持导出 Excel 与定时订阅。报表中心入口详见第 12.1 节。")
    page_break(doc)


# ----------- 第八章 -----------
def chapter8(doc):
    add_h1(doc, "第八章  库存业务管理")
    add_h2(doc, "8.1  库存查询")
    add_para(doc, "进入“库存业务 → 库存查询”，可按产品关键词与仓库两个条件检索实时库存，"
                  "并可通过表格列头的漏斗/排序图标对产品编码、产品名称、批次号、序列号、库存状态、数量、到期日等字段进一步筛选与排序。"
                  "工具栏右侧提供“导出”和橙色“库存调整”按钮。")
    add_image(doc, "15_inventory.png", caption="图 8-1  库存查询")
    add_table_kv(doc,
        ["列名", "说明"],
        [
            ["产品ID / 产品编码 / 产品名称", "库存对应产品的主键、编码（如 PROD-000195）与名称（如“演示产品-195”）。"],
            ["仓库ID / 仓库", "库存所在仓库的主键与名称，截图中均为“一级经销商-1-主仓”。"],
            ["批次号 / 序列号", "批号产品显示批号（如 B20260619-046），序列号产品显示单件序列号；未启用序列号时显示“-”。"],
            ["库存状态", "U 合格（可售）、B 不合格（冻结）、Q 待检（待质检放行），以彩色单字母标签显示。"],
            ["数量", "该批次/序列号在该仓库的实际结存数量。"],
            ["到期日", "产品失效日期，格式 YYYY-MM-DD；近效期与过期批次按颜色提醒。"],
            ["入库来源", "MIGRATION 表示历史数据迁移、PURCHASE 表示采购入库、SALES_RETURN 表示销退入库等。"],
            ["操作", "蓝色“查看”按钮打开库存详情抽屉，展示该批次的事务流水。"],
        ],
        col_widths=[4.0, 11.0])
    add_h2(doc, "8.2  库存移动")
    add_para(doc, "进入“库存业务 → 库存移动”，用于库位间调拨或仓库间调拨。"
                  "新建移动单时选择源仓库/货位与目标仓库/货位，添加产品与数量，提交后仓库复核确认，库存按源减目标增。")
    add_h2(doc, "8.3  库存调整")
    add_para(doc, "进入“库存业务 → 库存调整”，用于处理报损、报溢、状态转换（如待检转合格、合格转不合格）。"
                  "调整单必须选择调整类型与原因，上传佐证图片，审批通过后生效。")
    add_h2(doc, "8.4  库存盘点")
    add_para(doc, "进入“库存业务 → 库存盘点”，按仓库或分类创建盘点单。系统生成账面库存表，"
                  "仓库实地盘点后录入实盘数量，系统自动计算盘盈/盘亏差异。差异审批通过后生成库存调整单。"
                  "支持移动端扫码盘点，详见第十三章。")
    add_h2(doc, "8.5  效期预警")
    add_para(doc, "进入“库存业务 → 效期预警”，按产品设置的“临期预警月”阈值，集中展示即将到期的库存批次，"
                  "数据表格列依次为 ID、产品名称、批号、到期日、仓库、数量、状态（如 NORMAL 正常）、预警级别（如 WARNING 警告）、"
                  "创建时间、更新时间与操作列。列表默认按到期日升序排列，顶部提供产品名称、批号、仓库、预警级别（正常/警告/严重）"
                  "四个查询条件及“查询、重置”按钮，右上角为“导出、批量处理”按钮，底部分页显示总条数、每页条数与页码。")
    add_image(doc, "29_expiry.png", caption="图 8-2  效期预警")
    add_bullet(doc, "筛选条件：产品名称、批号、仓库、预警级别（正常 / 警告 / 严重）。")
    add_bullet(doc, "可直接在列表中对近效期批次发起促销或锁定，避免流入销售订单。")
    add_bullet(doc, "支持导出近效期清单并订阅邮件/企业微信通知。")
    add_h2(doc, "8.6  库存报表")
    add_para(doc, "在“数据看板 → 报表中心”选择“库存周转”报表，页面顶部提供产品分类筛选及“查询、重置、保存视图”按钮，"
                  "右上角提供“表格、图表、图表+表格、刷新、导出 xlsx”切换。KPI 卡片显示总库存、合格库存、待检库存、"
                  "不合格库存四个指标（如总库存 259,778、合格库存 235,080、待检库存 15,556、不合格库存 9,142），"
                  "柱状图按产品展示库存分布，表格列依次为编码、产品、规格、分类、当前库存、合格、待检、不合格、"
                  "近 30 天入、近 30 天出与平均周转，帮助识别滞销与积压。")
    add_image(doc, "31_inventory_report.png", caption="图 8-3  库存周转报表")
    page_break(doc)


# ----------- 第九章 -----------
def chapter9(doc):
    add_h1(doc, "第九章  产品追溯与序列号管理")
    add_h2(doc, "9.1  序列号与批次追溯")
    add_para(doc, "进入“库存业务 → 序列号追溯”，页面顶部提供“序列号追溯”与“批次追溯”两个切换标签，"
                  "右侧为搜索输入框与蓝色“查询”按钮。在“批次追溯”标签下输入批号（如 BATCH-DEMO-001）并点击“查询”，"
                  "下方“批次信息”区以两列栅格展示批号与事件数；事件表格列依次为时间、事件、序列号、数量、仓库，"
                  "若该批号暂无流转事件则显示“暂无数据”占位；输入真实业务批号后，表格会按时间正序展示该批次的全部库存事务"
                  "（采购入库 → 库存移动 → 销售出库 → 手术报台等），每行记录时间、事件类型、序列号、数量与仓库，"
                  "实现来源可查、去向可追。")
    add_image(doc, "16_trace.png", caption="图 9-1  序列号追溯")
    add_image(doc, "42_compliance.png", caption="图 9-2  批次追溯")
    add_h2(doc, "9.2  UDI 合规追溯")
    add_para(doc, "对于启用了“UDI 追溯”的产品，系统在采购收货、销售出库、库存移动等环节强制采集 UDI 编码（DI+PI），"
                  "并保留原始条码图片。UDI 数据可在追溯页导出，用于向国家医疗器械唯一标识数据库上报或接受监管检查。")
    add_h2(doc, "9.3  订单追溯报表")
    add_para(doc, "进入“数据看板 → 报表中心 → 订单追溯”，可按时间范围（近 30 天 / 本年等快捷选项 + 起止日期）、"
                  "订单状态、订单类型筛选订单的完整链路。页面顶部提供“查询、重置、保存视图”按钮，"
                  "右上角提供“表格、图表、图表+表格、刷新、导出 xlsx”切换。KPI 卡片显示订单数与订单总额"
                  "（如近 30 天订单数 100、订单总额 ¥ 1,325,751.34），柱状图按订单号展示订单金额分布，"
                  "表格列依次为订单号、类型、经销商、订单状态、订单金额、产品数、下单时间、审核时间、出库时间等，"
                  "覆盖订单审批、出库、收货全链路节点，支持导出追溯报告。")
    add_image(doc, "23_trace_report.png", caption="图 9-3  订单追溯报表")
    add_note(doc, "追溯数据采用“只增不改不删”的原则，任何修正都通过红冲单据完成，以保证审计链完整。")
    page_break(doc)


# ----------- 第十章 -----------
def chapter10(doc):
    add_h1(doc, "第十章  手术报告与市场营销")
    add_h2(doc, "10.1  手术报告")
    add_para(doc, "进入“手术与营销 → 手术植入报台”，维护植入/介入手术的跟台使用记录。"
                  "列表顶部提供关键词、状态、开始日期、结束日期四个查询条件及“查询/重置”按钮，"
                  "数据表格列依次为编号、报台单号、经销商、医院、手术日期、主刀医生、附件、状态、创建时间、更新时间与操作列，"
                  "右上角提供“新增”按钮，暂无数据时居中显示“暂无数据”。")
    add_image(doc, "17_surgery.png", caption="图 10-1  手术植入报台列表")
    add_para(doc, "新建手术报告时需录入经销商、医院、仓库（扣减合格库存）、手术日期、患者姓名、主刀医生、备注与现场照片，"
                  "并在“产品明细”区通过“添加产品”录入使用产品及序列号/UDI、术中用量，提交后系统自动扣减合格库存并保留报台记录。")
    add_h2(doc, "10.2  促销活动")
    add_para(doc, "进入“手术与营销 → 促销活动”，维护面向经销商或医院的促销政策，如买赠、满减、阶梯返点等。"
                  "促销单包含活动名称、活动时间、适用产品/经销商、促销规则、预算金额，审批通过后在销售下单时自动套用。")
    add_image(doc, "18_promotion.png", caption="图 10-2  促销活动")
    add_h2(doc, "10.3  移动端手术报台")
    add_para(doc, "跟台销售员通常在术后第一时间通过移动端录入手术报告，详见第十三章 13.5 节。移动端报台支持现场拍照上传、"
                  "扫码录入产品序列号、离线暂存，网络恢复后自动同步。")
    page_break(doc)


# ----------- 第十一章 -----------
def chapter11(doc):
    add_h1(doc, "第十一章  审批中心与工作流")
    add_h2(doc, "11.1  审批流配置")
    add_para(doc, "进入“审批中心 → 审批流配置”，管理员可针对采购订单、销售退货、采购退货、合同、授权等业务类型"
                  "分别维护审批模板。页面顶部提供业务类型、状态、名称/编码三个筛选条件与“查询”按钮，"
                  "右上角为蓝色“新建审批流”按钮。表格列依次为：ID、名称、业务类型、版本、优先级、状态、驳回策略、更新时间、操作。")
    add_image(doc, "19_workflow.png", caption="图 11-1  审批流配置列表")
    add_para(doc, "系统预置 5 套默认模板（截图中所有模板状态均为“已启用”，版本 1，优先级 10，驳回策略均为"
                  "“退回发起人修改后重新提交”）：采购订单默认审批模板、销售退货默认审批模板、采购退货默认审批模板、"
                  "合同默认审批模板、授权默认审批模板。操作列提供“编辑、停用、新版本”三个按钮——"
                  "“新版本”用于在保留历史版本的前提下发布新流程，旧流程实例仍按原版本走完。")
    add_table_kv(doc,
        ["元素", "说明"],
        [
            ["业务类型", "采购订单 / 采购退货 / 销售订单 / 销售退货 / 合同 / 授权 / 费用 等。"],
            ["版本与优先级", "同一业务类型可有多个版本，优先级数值越大越优先；停用的模板不参与流程匹配。"],
            ["发起人节点", "确定哪些角色/部门可以发起该流程。"],
            ["审批节点", "设置单人审批、多人会签、按比例通过等策略，可指定审批人或按角色/上级自动解析。"],
            ["条件分支", "根据单据字段（如金额 > 10 万、经销商级别）自动选择不同分支。"],
            ["驳回策略", "“退回发起人修改后重新提交”或“退回上一节点”。"],
            ["抄送人", "审批结果以消息/邮件形式抄送给相关人员。"],
            ["超时策略", "超过 N 小时未处理自动提醒，超过 M 小时自动升级或转交。"],
        ],
        col_widths=[3.0, 12.0])
    add_h2(doc, "11.2  我的审批")
    add_para(doc, "进入“审批中心 → 我的审批”，默认展示当前用户“待我审批”的任务，切换标签可查看“我已审批”"
                  "“我发起的”“抄送我的”。待办列表显示业务类型、单号、发起人、提交时间、当前节点。")
    add_image(doc, "20_approval.png", caption="图 11-2  我的审批")
    add_para(doc, "处理审批操作：")
    add_number(doc, "点击“查看”打开业务单据详情，核对信息与附件。")
    add_number(doc, "点击“同意”/“驳回”按钮，在弹出框中填写审批意见。")
    add_number(doc, "同意后流程进入下一节点；驳回后流程回到发起人，发起人可修改后重新提交。")
    add_number(doc, "支持“转交”（将审批任务转给他人）与“加签”（临时增加审批人）。")
    add_h2(doc, "11.3  审批监控")
    add_para(doc, "进入“审批中心 → 审批监控”，管理员可查看所有流程实例的当前状态、耗时、是否超时。"
                  "对超时任务可执行催办、强制结束、跳转节点等操作。")
    add_image(doc, "30_approval_monitor.png", caption="图 11-3  审批监控")
    page_break(doc)


# ----------- 第十二章 -----------
def chapter12(doc):
    add_h1(doc, "第十二章  数据看板与业务报表")
    add_h2(doc, "12.1  报表中心")
    add_para(doc, "进入“数据看板 → 报表中心”，页面以卡片网格形式列出全部预置报表，右上角标注版本号（如 v4.2）。"
                  "报表按业务分类组织：销售类（销售业绩排行、产品销售 TOP10、销售明细）、库存类（库存周转、库存呆滞/超期）、"
                  "订单类（订单追溯、拒单率/审批时长）、财务类（应收款项、返利/折扣对账）、合同/授权类（合同台账、授权余额/超期、"
                  "借货余额/超期）、报台与画像类（手术报台统计）。每张卡片显示报表名称、功能简介、时间范围、筛选条件数与返回列数，"
                  "点击卡片进入对应报表详情页。所有报表均提供时间区间、经销商、产品分类等筛选条件，支持明细/汇总切换、"
                  "图表/表格切换、导出 xlsx 与订阅。")
    add_image(doc, "43_purchase_report.png", caption="图 12-1  报表中心")
    add_h2(doc, "12.2  销售业绩排行报表")
    add_para(doc, "销售业绩排行报表按经销商维度展示销售金额、订单数、客单价、活跃经销商数等指标。"
                  "页面顶部筛选区提供时间范围（本年/本月/近 30 天等快捷选项 + 起止日期）、经销商级别、区域（模糊）、"
                  "订单状态、订单类型五个筛选条件及“查询、重置、保存视图”按钮，右上角提供“表格、图表、图表+表格、刷新、导出 xlsx”切换。"
                  "图表区以蓝色柱状图展示各经销商销售额排行，表格列依次为编码、经销商、级别、区域、订单数、销售总额、客单价、"
                  "通过、草稿、取消、最近下单、首次下单。KPI 卡片显示总销售额（如 ¥ 8,705,187.46）、订单数（如 695）、"
                  "平均客单价（如 ¥ 12,520.86）、活跃经销商数（如 50）。")
    add_image(doc, "32_sales_report.png", caption="图 12-2  销售业绩排行报表")
    add_h2(doc, "12.3  采购报表")
    add_para(doc, "采购报表按供应商、产品统计采购数量、金额、到货及时率，用于评估供应商表现与采购计划，可在报表中心对应卡片进入。")
    add_h2(doc, "12.4  库存周转报表")
    add_para(doc, "库存周转报表展示各产品在统计期内的总库存、合格库存、待检库存、不合格库存以及近 30 天入/出库数量，"
                  "帮助识别滞销与积压。页面顶部提供产品分类筛选及“查询、重置、保存视图”按钮，"
                  "KPI 卡片显示总库存、合格库存、待检库存、不合格库存四个指标，柱状图按产品展示库存分布，"
                  "表格列依次为编码、产品、规格、分类、当前库存、合格、待检、不合格、近 30 天入、近 30 天出、平均周转。参见图 8-3。")
    add_h2(doc, "12.5  报表订阅")
    add_para(doc, "在“业务报表 → 报表订阅”中统一管理定时推送任务，页面右上角为“新建订阅”按钮，"
                  "数据表格列依次为 ID、订阅名称、报表、频率、收件人、状态、上次运行、结果与操作。"
                  "新建订阅时需选择目标报表、设置订阅周期（每日 / 每周 / 每月）、收件人与文件格式（Excel/PDF），"
                  "系统在指定时间将报表推送到收件人的消息中心与邮箱；暂无订阅时列表显示“暂无订阅，点击右上角新建”。")
    add_image(doc, "40_message_config.png", caption="图 12-3  报表订阅管理")
    page_break(doc)


# ----------- 第十三章 移动端 -----------
def chapter13(doc):
    add_h1(doc, "第十三章  移动端 H5 操作说明")
    add_para(doc, "移动端 H5 面向经常外出的销售员、跟台工程师、仓库巡查员与需要随时审批的管理者，"
                  "采用 Vue 3 + Vant 4 构建，支持 iOS 与 Android 主流浏览器及企业微信/钉钉/飞书内置浏览器。"
                  "登录后底部固定 Tab 栏包含：首页、订单、报台、审批、消息、我的。")

    add_h2(doc, "13.1  移动首页")
    add_para(doc, "移动首页采用蓝色渐变顶栏，左上角显示 DMS logo 与“你好，”问候语，右上角显示当前日期（如 2026/8/16）。"
                  "页面自上而下分为三个区域：")
    add_image(doc, "m_home.png", caption="图 13-1  移动首页")
    add_table_kv(doc,
        ["区域", "内容"],
        [
            ["今日业绩", "展示“今日销售金额”（如 ¥ 32,533.43）与“今日订单数”（如 3），数据实时刷新。"],
            ["本月业绩", "展示“本月销售金额”（如 ¥ 1,607,687.04）、“本月订单数”（如 123）、“本月报台数”（如 30），"
                       "并在右侧提供“查看趋势 › 业绩详情”入口，点击进入“我的业绩”图表页。"],
            ["快捷入口", "两行四列图标宫格：下销售订单（蓝色）、填手术报台（绿色）、我的订单、我的业绩、"
                       "扫码收货、库存扫码、库存盘点、移动审批，单击直达对应功能。"],
            ["最近订单", "按时间倒序展示最近订单卡片：订单号（如 SO-00000001）、经销商、金额（红色，如 ¥10953.64）、状态标签，"
                       "点击卡片进入订单详情。"],
        ],
        col_widths=[2.6, 12.4])
    add_para(doc, "底部固定 Tab 栏自左向右依次为：首页、订单、报台、审批、消息、我的，"
                  "当前 Tab 以蓝色高亮，其他 Tab 为灰色线框图标。")

    add_h2(doc, "13.2  销售订单（移动端）")
    add_para(doc, "点击底部“订单”Tab 或首页“我的订单”快捷入口，进入销售订单列表。"
                  "页面顶部蓝色标题栏显示“销售订单”，右上角有“+”新增按钮；下方为“搜索单号/经销商”搜索框。"
                  "列表按白色卡片形式逐张展示订单，每张卡片包含：订单号（如 SO-00000001）、"
                  "经销商名称（如苏州康宁医疗器械有限公司）、订单日期（如 2026-08-15）、"
                  "红色金额（如 ¥ 10953.64）以及橙色“草稿”状态标签。"
                  "右下角悬浮蓝色“+ 下销售订单”按钮，可在任意滚动位置快速下单。")
    add_image(doc, "m_orders.png", caption="图 13-2  移动销售订单列表")
    add_h3(doc, "13.2.1  下销售订单")
    add_para(doc, "点击“+ 下销售订单”悬浮按钮，进入“下销售订单”页。页面自上而下分为三个区块：")
    add_image(doc, "m_order_create.png", caption="图 13-3  移动端下销售订单")
    add_table_kv(doc,
        ["区块", "字段与操作"],
        [
            ["基本信息", "经销商（必填，红色星号，点击选择“您负责的”经销商）、仓库（选择发货仓库，可选）、"
                       "订单类型（默认为“常规”，点击切换常规/紧急/缺货/定制）、订单日期（默认当天 2026-08-16）、备注（选填）。"],
            ["产品明细", "在未选择经销商时产品行显示灰色提示“请先选择经销商”；选定经销商后点击“+ 添加产品”按钮，"
                       "可按产品编码搜索或扫码添加产品，逐行录入数量、单价；每行自动计算小计。"],
            ["金额合计", "实时汇总产品数量（0 项）、不含税金额（¥ 0.00）、税额（¥ 0.00），"
                       "含税总额随明细自动计算。底部固定蓝色“提交订单”按钮，点击后校验必填项并提交。"],
        ],
        col_widths=[2.8, 12.2])
    add_para(doc, "移动端针对单手操作进行了优化：必填字段以红色星号标注，数字输入框自动弹出数字键盘，"
                  "点击顶部左上角“‹”返回可放弃当前录入。")
    add_h3(doc, "13.2.2  订单详情与审批")
    add_para(doc, "点击订单卡片进入订单详情，可查看订单行、金额、审批进度；若当前用户是审批人，页面底部显示“同意/驳回”按钮。")

    add_h2(doc, "13.3  移动审批")
    add_para(doc, "点击底部“审批”Tab 进入移动审批列表，按“待办/已办/我发起”标签切换。"
                  "待办列表显示业务类型、单号、发起人、提交时间；点击进入审批详情，可查看单据内容与附件，"
                  "底部提供“同意”“驳回”“转交”按钮。")
    add_image(doc, "27_mobile_approval.png", caption="图 13-4  移动审批")
    add_note(doc, "移动审批支持指纹/面容解锁后快捷审批（需在“我的 → 设置”中开启生物识别）。")

    add_h2(doc, "13.4  扫码收货与库存查询")
    add_para(doc, "在首页点击“扫码收货”进入扫码收货页。页面顶部为摄像头预览区，下方提供“打开摄像头扫描”与“手动输入”两个按钮；"
                  "扫描产品条码/UDI 后系统自动带入产品信息，操作员录入数量、批号并提交，生成收货记录。")
    add_image(doc, "26_mobile_scan.png", caption="图 13-5  扫码收货")
    add_para(doc, "在首页点击“库存扫码”进入库存扫码查询页，扫描产品条码或输入批号/序列号后，"
                  "页面显示该产品在各仓库的库存数量、批号、效期与状态，方便仓库巡查与跟台备货。")
    add_image(doc, "35_mobile_stock.png", caption="图 13-6  库存扫码查询")

    add_h2(doc, "13.5  手术报台（移动端）")
    add_para(doc, "点击底部“报台”Tab 进入手术报告列表。页面顶部蓝色标题栏显示“手术报台”及绿色手术刀图标，"
                  "右上角“+”为新增按钮。当前账号下暂无报台记录时，列表区显示“没有更多了”、"
                  "绿色空状态图标与“暂无报台，点击右下角新建”提示文案，右下角悬浮蓝色“+ 新建报台”按钮。")
    add_image(doc, "m_surgery.png", caption="图 13-7  手术报台列表（空状态）")
    add_para(doc, "点击“+ 新建报台”进入“手术植入报台”录入页，顶部为返回箭头与标题，下方分为“基本信息”和“产品明细”两个分组。"
                  "“基本信息”分组字段包括：经销商（必填，选择当前销售员负责的经销商）、医院（必填，选择经销商后才可选择）、"
                  "仓库（必填，用于扣减合格库存）、手术日期（必填，默认当天）、患者姓名（必填）、主刀医生、备注与现场照片（相机图标拍照上传）；"
                  "“产品明细”分组初始显示一行空明细，点击“+ 添加产品”可继续添加产品，未选经销商时产品行提示“请先选择经销商”；"
                  "页面底部为蓝色“提交报台”按钮，支持弱网暂存、拍照上传，网络恢复后自动同步。")
    add_image(doc, "m_surgery_create.png", caption="图 13-8  移动端手术报台录入")

    add_h2(doc, "13.6  我的业绩与消息")
    add_para(doc, "在首页“本月业绩”卡右侧点击“查看趋势 ›”进入“我的业绩”页面。页面顶部展示“本月销售”概览卡："
                  "左侧为蓝色“¥ 1,607,687.04”销售金额，右侧为蓝色“123”销售订单数；"
                  "下方为“近 12 月销售趋势”列表，从 2025-09 至 2026-08 逐月以蓝色横向进度条 + 红色金额的形式展示（如 2026-08 ¥1,607,687.04）；"
                  "最下方为“本月 TOP 经销商”榜单，红色序号徽章 + 经销商名称 + 金额（如第 1 名 北京京医通达有限公司 ¥89,936.63）。")
    add_image(doc, "m_dashboard.png", caption="图 13-9  我的业绩")
    add_para(doc, "点击底部“消息”Tab 可查看系统通知、审批提醒、预警消息，消息按业务类型分类，未读消息以红点提示。")
    add_image(doc, "m_messages.png", caption="图 13-10  移动消息中心")

    add_h2(doc, "13.7  我的")
    add_para(doc, "点击底部“我的”Tab 进入个人中心。顶部为蓝色渐变区，圆形头像图标下方显示当前用户姓名（如“林管理员”）"
                  "与所属组织类型（如“厂商”）。下方白色卡片依次展示“消息中心”与“我的审批”两个功能入口，右侧带“›”箭头；"
                  "再下方为基础信息卡，以两列形式展示：账号（sys_admin）、角色（系统管理员）、经销商（厂商账号无绑定经销商时为“-”）、"
                  "手机（未维护时为“-”）、邮箱（vinkinyu@163.com）。卡片底部为红色边框的“退出登录”按钮。")
    add_image(doc, "m_profile.png", caption="图 13-11  我的")
    page_break(doc)


# ----------- 第十四章 -----------
def chapter14(doc):
    add_h1(doc, "第十四章  系统管理与配置")
    add_h2(doc, "14.1  角色权限")
    add_para(doc, "进入“用户与权限 → 角色权限”，维护系统角色及其权限。页面顶部提供“角色名称 / 编码”查询条件及"
                  "“查询、重置”按钮，右上角为“新增角色”按钮。数据表格列依次为编号、角色编码（如 SYS_ADMIN、SALES_MGR、"
                  "SALES、CS、BIZ、FIN、CONTRACT_SPEC、DEALER_ADMIN 等）、角色名称（系统管理员、销售经理、销售、客服、商务、"
                  "财务、合同专员、经销商管理员等）、类型（自定义角色）、描述、状态（绿色“启用”标签）与操作列，"
                  "操作列提供“编辑”和蓝色“权限设置”按钮。权限粒度包括菜单权限（是否可见）、按钮权限（是否可点击"
                  "新增/编辑/删除/导出等按钮）、数据权限（本人 / 本部门 / 本部门及下级 / 全部）。")
    add_image(doc, "41_permission.png", caption="图 14-1  角色权限")
    add_para(doc, "点击“权限设置”打开权限分配抽屉，按菜单树勾选节点并在右侧勾选按钮级权限，"
                  "权限保存后实时生效，被授权用户重新登录或刷新页面即可看到最新菜单与按钮。")
    add_h2(doc, "14.2  数据字典")
    add_para(doc, "数据字典维护系统中使用的枚举值，如产品类型、订单状态、库存状态、审批结果、促销类型等。"
                  "字典按“字典分类 + 字典项”两级组织，分类对应一个下拉选项集合（如订单状态包含草稿、待审批、审批中、已通过、已驳回），"
                  "字典项包含编码、显示名称、排序值、是否启用。修改字典项后，相关页面的下拉选项实时刷新，"
                  "无需重新发版；停用的字典项在新单据中不可选，但历史单据仍保留原值显示以保证数据可追溯。")
    add_h2(doc, "14.3  列表页配置")
    add_para(doc, "进入“用户与权限 → 列表页配置”，页面顶部提供“页面名称 / 编码”查询条件及“查询、重置”按钮，右上角为“新增配置”按钮。"
                  "数据表格列依次为页面编码（如 products、orders、sales-outs、dealers、purchase-orders 等）、页面名称、列表标题、备注、"
                  "创建时间、更新时间与操作列（编辑、删除）。管理员可按角色/岗位配置每个列表页的搜索字段、工具栏按钮、行内按钮，"
                  "例如可配置“销售岗位”只看到“关键词、状态、开始日期、结束日期、经销商”五个搜索项，并隐藏“批量删除”等敏感按钮；"
                  "也可调整列的显示顺序与默认宽度。配置项以 JSON 存储，按“页面编码 + 角色编码”作为唯一键，"
                  "用户登录后自动加载对应配置，无需重启服务。")
    add_image(doc, "45_param.png", caption="图 14-2  列表页配置")

    add_h2(doc, "14.4  系统参数与导入导出任务")
    add_para(doc, "系统参数提供全局开关与默认值配置，例如订单编号前缀、库存预警阈值、近效期月数、审批超时小时数、"
                  "移动端单次最大扫码数量等。导入导出任务以异步方式执行，长任务在后台处理，"
                  "完成后通过站内消息与邮件通知用户，可在“导入导出任务”页下载结果文件并查看失败原因。")
    add_table_kv(doc,
        ["参数分组", "说明"],
        [
            ["单据编号", "订单、出库、入库、退货等单据的前缀与流水位数"],
            ["库存预警", "合格库存下限、近效期预警月数、呆滞品判定天数"],
            ["审批流程", "审批超时小时数、超时自动转交、催办间隔"],
            ["移动端", "单次扫码上限、离线缓存天数、照片压缩质量"],
            ["消息通知", "站内信、邮件、短信开关及模板编号"],
        ])

    add_h2(doc, "14.5  租户与菜单管理")
    add_para(doc, "平台管理员可在租户管理中维护租户信息、到期时间、用户数上限、模块授权；"
                  "菜单管理支持多级菜单树，可配置菜单名称、路径、图标、排序、是否在移动端显示等。"
                  "所有变更均记录操作日志，可在“系统管理 → 操作日志”中按人、时间、模块检索。")
    page_break(doc)


# ----------- 第十五章 -----------
def chapter15(doc):
    add_h1(doc, "第十五章  消息中心与日志审计")
    add_h2(doc, "15.1  消息中心")
    add_para(doc, "PC 端右上角铃铛图标进入消息中心，页面顶部提供“全部、未读、已读”三个状态切换 Tab、"
                  "“分类”下拉筛选与“全部已读”按钮；数据表格列依次为状态、标题、内容、分类与时间。"
                  "消息类型包括系统通知、审批提醒、库存预警、近效期提醒、订单状态变更、导入导出结果等。"
                  "未读消息在状态列以圆点标识，点击消息标题可跳转到对应业务单据，处理后消息自动置为已读。")
    add_image(doc, "33_message.png", caption="图 15-1  消息中心")

    add_h2(doc, "15.2  消息模板与订阅")
    add_para(doc, "管理员可在“消息模板”中维护站内信、邮件、短信三种渠道的模板内容，"
                  "模板使用 ${变量名} 占位，发送时由业务事件自动填充。"
                  "“报表订阅”允许用户订阅每日/每周库存与销售报表，系统按 cron 定时生成并推送到邮箱。")

    add_h2(doc, "15.3  操作日志与登录日志")
    add_para(doc, "系统对所有新增、修改、删除、审批、导出、登录等操作记录审计日志，"
                  "包括操作人、IP、操作时间、模块、操作类型、请求参数、操作结果。"
                  "日志列表支持按模块、操作人、时间、结果检索，并可导出 Excel。"
                  "业务单据详情页还可查看“单据流转日志”，呈现每次状态变化、审批意见与附件。"
                  "在“日志与监控 → 登录日志”中可单独查询用户登录记录，"
                  "列表列依次为用户名、姓名、IP、位置、浏览器、操作系统、状态与操作时间，"
                  "顶部提供用户名、IP、状态、开始日期、结束日期五个查询条件及“查询、重置”按钮，"
                  "右上角提供“导出”按钮，方便安全审计与异常登录追溯。")
    add_image(doc, "24_log.png", caption="图 15-2  登录日志")

    add_h2(doc, "15.4  邮件发送日志")
    add_para(doc, "所有由系统触发的邮件均记录发送日志，包含收件人、主题、模板、发送时间、发送状态、"
                  "SMTP 返回信息。发送失败的邮件支持手动重发，便于排查 SMTP 配置与收件箱过滤问题。")
    page_break(doc)


# ----------- 第十六章 -----------
def chapter16(doc):
    add_h1(doc, "第十六章  常见问题与故障排查")
    add_h2(doc, "16.1  登录类问题")
    add_table_kv(doc,
        ["现象", "排查与解决"],
        [
            ["提示租户编码错误", "确认登录页“企业/租户编码”是否为 default 或管理员分配的编码"],
            ["提示账号或密码错误", "检查大小写、是否误开 Caps Lock；忘记密码由管理员重置"],
            ["登录后页面空白", "清除浏览器缓存后使用 Ctrl+Shift+R 强制刷新；建议使用 Chrome 120+"],
            ["手机提示会话过期", "在“我的”中重新登录；若长期不用，自动登出可保证账号安全"],
        ])

    add_h2(doc, "16.2  单据与审批")
    add_table_kv(doc,
        ["现象", "排查与解决"],
        [
            ["提交订单提示库存不足", "查看库存总览中合格库存可用量；若有在途入库，可先保存草稿"],
            ["审批人看不到待办", "检查审批流程模板是否启用、节点审批人是否含当前角色"],
            ["订单无法删除", "已审核、已出库或被退款单引用的订单不能删除，请使用红冲处理"],
            ["出库扫码提示序列号已使用", "该序列号已在其他出库单使用，请核对实物或在追溯中查询"],
        ])

    add_h2(doc, "16.3  库存与追溯")
    add_table_kv(doc,
        ["现象", "排查与解决"],
        [
            ["库存数量与实际不一致", "执行库存盘点，按账存生成差异，审批后调账；同时检查未审核出入库单"],
            ["近效期未预警", "确认系统参数“近效期预警月数”是否正确；消息订阅是否勾选"],
            ["追溯无数据", "检查产品是否启用序列号/批号管理；stock_serials 是否有入库记录"],
            ["报表数据延迟", "报表基于定时聚合表，默认 10 分钟刷新一次；可手动触发刷新"],
        ])

    add_h2(doc, "16.4  移动端常见问题")
    add_table_kv(doc,
        ["现象", "排查与解决"],
        [
            ["扫码无反应", "授予相机权限；擦拭摄像头；在光线充足处重试；支持手动输入序列号"],
            ["离线提交失败", "网络恢复后在“我的 → 离线队列”中重试；超过 7 天的离线数据会被清理"],
            ["照片上传失败", "检查网络与图片大小，系统会自动压缩到 2MB 以内"],
            ["收不到推送", "检查手机系统设置中应用通知权限，并在“我的 → 消息设置”中开启对应类别"],
        ])

    add_h2(doc, "16.5  浏览器与环境要求")
    add_para(doc, "PC 端推荐 Chrome 120+、Edge 120+、360 极速浏览器 13+，屏幕分辨率 1440×900 及以上；"
                  "移动端支持 iOS 14+（Safari）、Android 9+（Chrome、微信内置浏览器）。"
                  "如遇页面布局错乱，请先强制刷新并清除浏览器缓存。")
    page_break(doc)


# ----------- 第十七章 -----------
def chapter17(doc):
    add_h1(doc, "第十七章  附录")
    add_h2(doc, "17.1  快捷键")
    add_table_kv(doc,
        ["快捷键", "功能"],
        [
            ["Alt + 1 / 2 / 3", "PC 端切换顶部主导航：工作台 / 业务 / 报表"],
            ["Ctrl + Enter", "弹窗中快速提交表单"],
            ["Esc", "关闭当前弹窗或抽屉"],
            ["Ctrl + Shift + R", "强制刷新浏览器并清除缓存"],
            ["F5", "刷新当前页"],
            ["↑ / ↓", "表格中上下移动选择行"],
        ])

    add_h2(doc, "17.2  术语表")
    add_table_kv(doc,
        ["术语", "含义"],
        [
            ["SKU", "库存单位，系统中以“产品 + 规格 + 注册证”唯一标识"],
            ["UDI", "医疗器械唯一标识，由 DI（产品标识）+ PI（生产标识）组成"],
            ["批号 / 序列号", "批号用于同批生产产品；序列号用于单件唯一追踪，高值耗材多用序列号"],
            ["合格库存", "经入库验收且质量状态为合格的可售库存"],
            ["近效期", "距离失效日期不足设定月数（默认 6 个月）的库存"],
            ["红冲", "对已审核单据生成反向单据以冲销原单据，保留审计痕迹"],
            ["报台", "手术使用登记，记录医院、术者、患者信息与耗材消耗明细"],
            ["授权", "厂商授予经销商在指定区域销售指定产品的授权书"],
            ["GSP", "医疗器械经营质量管理规范，系统流程按 GSP 要求设计"],
        ])

    add_h2(doc, "17.3  联系方式与版本记录")
    add_para(doc, "技术支持：请联系本单位系统管理员或 DMS 实施团队；如发现系统缺陷，请在“消息中心 → 新建反馈”中提交截图与复现步骤。")
    add_table_kv(doc,
        ["版本", "日期", "主要变更"],
        [
            ["V3.7.1", "2026-07-25", "完善移动端扫码、手术报台、审批流；新增防回归校验与端到端部署铁律"],
            ["V3.7.0", "2026-06-30", "新增报表中心、库存周转分析、合同工作台、消息订阅"],
            ["V3.6.0", "2026-05-20", "支持多租户、角色权限、列表页配置、操作日志"],
            ["V3.0.0", "2026-01-15", "首次发布，覆盖产品、订单、库存、追溯核心业务"],
        ])
    add_para(doc, "——本手册随系统迭代持续更新，最新版本以系统内置说明为准——")


# ----------- 主入口 -----------
def main():
    doc = Document()
    setup_page(doc.sections[0])
    add_header_footer(doc.sections[0])
    build_cover(doc)
    build_toc(doc)
    chapter1(doc)
    chapter2(doc)
    chapter3(doc)
    chapter4(doc)
    chapter5(doc)
    chapter6(doc)
    chapter7(doc)
    chapter8(doc)
    chapter9(doc)
    chapter10(doc)
    chapter11(doc)
    chapter12(doc)
    chapter13(doc)
    chapter14(doc)
    chapter15(doc)
    chapter16(doc)
    chapter17(doc)
    doc.save(OUTPUT_DOCX)
    print(f"OK saved: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()