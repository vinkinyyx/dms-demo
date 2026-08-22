# -*- coding: utf-8 -*-
"""估算 docx 页数（通过段落数+图片占位+表格粗略估算，准确值需在 Word 中查看）"""
import os
from docx import Document

files = {
    "申请表": r"d:\Workspace\TRAE\DMS\软著申请材料\01_申请表\计算机软件著作权登记申请表_DMS经销商管理系统V1.0.docx",
    "源代码": r"d:\Workspace\TRAE\DMS\软著申请材料\02_源代码\源程序_DMS经销商管理系统V1.0_前30页后30页.docx",
    "用户手册": r"d:\Workspace\TRAE\DMS\软著申请材料\03_用户手册\DMS经销商管理系统V1.0用户操作手册.docx",
}

for name, path in files.items():
    doc = Document(path)
    paras = len(doc.paragraphs)
    tables = len(doc.tables)
    # 粗略估算：正文一页约 35 段（1.5 倍行距），表格平均 8 行
    page_est = paras / 35 + tables * 0.3
    print(f"{name}: 段落 {paras}, 表格 {tables}, 估算页数 ≈ {page_est:.1f}")
