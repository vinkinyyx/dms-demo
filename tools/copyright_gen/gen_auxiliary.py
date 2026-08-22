# -*- coding: utf-8 -*-
"""
生成软著申请的附属材料：
1. 作品说明书（简要说明软件功能与开发情况）
2. 原创版本声明（如版本号高于 V1.0 时使用；本项目是 V1.0 可省略，作为备选）
3. 源代码示例说明（页眉/目录/注释样例指引）
4. 申请人身份证明模板（企业营业执照副本复印件/个人身份证复印件粘贴页）
5. 委托书模板（委托代理机构时使用）
6. 提交材料清单与注意事项
"""
import os
from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = r"d:\Workspace\TRAE\DMS\软著申请材料\04_其他证明材料"
os.makedirs(OUT_DIR, exist_ok=True)


def set_font(run, name=u'宋体', size=10.5, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)


def set_section(section):
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Mm(25)
    section.bottom_margin = Mm(25)
    section.left_margin = Mm(28)
    section.right_margin = Mm(25)


def title(doc, text, size=18):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(text), name=u'黑体', size=size, bold=True)


def h(doc, text, size=12):
    p = doc.add_paragraph()
    set_font(p.add_run(text), name=u'黑体', size=size, bold=True)


def para(doc, text, indent=True, size=10.5):
    p = doc.add_paragraph()
    if indent:
        pPr = p._p.get_or_add_pPr()
        ind = OxmlElement('w:ind'); ind.set(qn('w:firstLineChars'), '200'); ind.set(qn('w:firstLine'), '420')
        pPr.append(ind)
    set_font(p.add_run(text), size=size)


def signature_block(doc):
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(p.add_run("申请人（盖章/签字）：____________________"), size=11)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(p.add_run("日　期：　　　 年　　 月　　 日"), size=11)


def make_function_intro():
    """作品/软件功能说明书（2-3 页）"""
    f = os.path.join(OUT_DIR, "附属材料1_软件功能说明书.docx")
    doc = Document(); set_section(doc.sections[0])
    style = doc.styles['Normal']; style.font.name = u'宋体'; style.font.size = Pt(10.5)

    title(doc, "软件功能与开发情况说明书")
    doc.add_paragraph()

    h(doc, "一、软件基本信息")
    para(doc, "软件全称：DMS经销商管理系统 V1.0")
    para(doc, "软件简称：DMS系统")
    para(doc, "版本号：V1.0")
    para(doc, "开发完成日期：2026年08月15日")
    para(doc, "首次发表日期：未发表")
    para(doc, "开发方式：单独开发")
    para(doc, "权利取得方式：原始取得")

    h(doc, "二、软件开发目的")
    para(doc,
         "本软件面向医疗器械与医用耗材生产厂家及分销企业，针对其在厂家主数据下发、经销商资质管理、"
         "合同与销售授权、销售与采购协同、多仓库存管控、植入类产品序列号追溯、手术报台、"
         "审批流程合规等环节存在的信息化薄弱问题，采用多租户 SaaS 模式，构建覆盖"
         "“厂家—经销商—医院终端”"
         "全链路的业务管理平台，帮助企业实现业务在线化、数据可追溯、流程合规化。")

    h(doc, "三、主要功能")
    bullets = [
        "主数据管理：产品、分类、产品线、经销商、医院终端、仓库、区域、价格政策、供应商等统一维护；",
        "合同与授权：合同模板、合同在线编制、审批流转、经销商授权范围与有效期管理、到期预警；",
        "销售管理：销售订单、销售出库、销退管理、销售价格政策、订单执行跟踪；",
        "采购管理：采购订单、收货入库、采退管理、供应商资质管理；",
        "仓储管理：库存查询、库存移动、库存调整、库存盘点、批次与序列号追溯、效期预警；",
        "手术管理：植入手术报台、产品序列号与患者使用记录关联、手术附件管理；",
        "促销管理：起订量、满额减、满赠、阶梯价等促销策略；",
        "审批工作流：可视化审批流配置，支持顺序审批、会签、或签、加签、转办、委托和超时催办；",
        "报表分析：数据驾驶舱、销售排行、产品 TOP10、库存周转、订单追溯、经销商 360 画像；",
        "平台管理：多租户、用户角色、菜单按钮、数据权限、数据字典、消息中心、日志审计；",
        "移动端 H5：移动下单、扫码收货/出库、手术报台、移动审批、消息提醒。",
    ]
    for b in bullets:
        para(doc, "● " + b, indent=False)

    h(doc, "四、技术特点")
    tech = [
        "采用前后端分离架构：后端基于 Spring Boot 3.2 + Java 17 + Spring Security + JPA/MyBatis-Plus；"
        "前端基于 Vue 3 + Vite 5 + Element Plus + Pinia；移动端采用 Vant 4。",
        "多租户行级隔离：基于 TenantContext 与 Hibernate 拦截器自动注入租户条件，支持厂家租户与经销商租户两级模型。",
        "自研审批流引擎：支持模板版本化、条件分支、顺序/会签/或签、前加签/后加签、转办委托、抄送和管理员改派终止。",
        "库存并发控制：采用乐观条件更新（UPDATE … WHERE qty >= ?）与数据库行级锁，保证批次与序列号库存的一致性。",
        "业务单号规则：采用 PREFIX-YYYYMMDD-NNNNN 规则，由数据库序列表原子自增生成，避免并发重复。",
        "审计日志：基于 AOP 切面自动采集操作日志，敏感字段脱敏，同时写入数据库与按日滚动日志文件。",
        "报表查询：采用 WITH CTE + LEFT JOIN 形式，避免相关子查询导致的数据库异常，提升大数据量下的查询稳定性。",
        "数据库版本化：使用 Flyway 管理数据库脚本，保障多环境、多版本升级的一致性。",
        "容器化部署：基于 Docker Compose 与 Nginx 实现容器化部署，支持测试、生产双环境隔离。",
        "统一接口规范：所有接口统一返回 {code, message, data} 结构，业务异常使用错误码枚举，便于排查和前端处理。",
    ]
    for t in tech:
        para(doc, "● " + t, indent=False)

    h(doc, "五、开发环境与运行环境")
    para(doc,
         "开发环境：Windows 10/11、JDK 17、Node.js 18、Maven 3.9、PostgreSQL 14、Redis 7、IntelliJ IDEA、Visual Studio Code。")
    para(doc,
         "服务端运行环境：CentOS 7.9/Ubuntu 22.04、JDK 17、PostgreSQL 14、Redis 7、Nginx 1.24、Docker 24；"
         "建议 4 核 CPU、8GB 内存、100GB 硬盘。")
    para(doc,
         "客户端运行环境：Chrome/Edge/Firefox 最新两个稳定版本；移动端支持 iOS 14+/Android 10+。")

    h(doc, "六、编程语言与源程序量")
    para(doc,
         "主要编程语言：Java、JavaScript、SQL、HTML、CSS；总源程序量约 65589 行，"
         "其中 Java 约 40330 行、前端 Vue/JavaScript 约 18176 行、SQL 脚本约 7083 行。")

    h(doc, "七、独创性声明")
    para(doc,
         "本软件由申请人自主研发，整体架构、数据库设计、业务模块、审批流引擎、报表查询逻辑、"
         "前端界面设计及源代码均为独立创作，未抄袭、复制他人享有著作权的作品，"
         "也未使用侵犯第三方知识产权的代码或素材。软件中使用的开源组件均遵循其相应开源许可协议，"
         "与申请人自研代码在工程结构上独立可辨。")

    signature_block(doc)
    doc.save(f)
    print("OK:", f)


def make_copyright_affidavit():
    """原创性/不侵权声明"""
    f = os.path.join(OUT_DIR, "附属材料2_原创性声明.docx")
    doc = Document(); set_section(doc.sections[0])

    title(doc, "原创性与不侵权声明")
    doc.add_paragraph()
    para(doc,
         "本人/本单位就所申请登记的“DMS经销商管理系统 V1.0”（以下简称“本软件”）"
         "作出如下声明：")
    items = [
        "本软件由本人/本单位独立开发完成，系原创作品，不存在抄袭、复制他人作品的情形；",
        "本软件所使用的算法、业务流程、界面设计、文档资料均为合法取得或自主创作；",
        "本软件中使用的开源组件遵循其开源许可协议，相关声明与许可文本已在软件发行包中予以保留；",
        "本软件不含有任何违反国家法律法规、损害社会公共利益或侵犯他人合法权益的内容；",
        "如因本软件著作权归属或侵权问题引发纠纷，由本人/本单位独立承担全部法律责任，"
        "与中国版权保护中心及登记机关无关。",
    ]
    for i, t in enumerate(items, 1):
        para(doc, f"{i}. {t}")
    doc.add_paragraph()
    para(doc, "特此声明。")
    signature_block(doc)
    doc.save(f)
    print("OK:", f)


def make_identity_template():
    """身份证明粘贴模板"""
    f = os.path.join(OUT_DIR, "附属材料3_身份证明粘贴模板.docx")
    doc = Document(); set_section(doc.sections[0])

    title(doc, "申请人身份证明")
    doc.add_paragraph()

    h(doc, "一、申请主体为企业法人的")
    para(doc, "请提交以下证明文件的复印件，并加盖企业公章：")
    items = [
        "企业法人营业执照副本（统一社会信用代码证书）复印件；",
        "如为事业单位、社会团体法人，提交事业单位法人证书/社会团体法人登记证书复印件；",
        "经办人身份证复印件（正反面）。",
    ]
    for it in items:
        para(doc, "● " + it, indent=False)

    # 营业执照粘贴框
    doc.add_paragraph()
    h(doc, "【营业执照（副本）复印件粘贴处】")
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    cell.width = Cm(15)
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("\n\n\n（请将营业执照副本复印件在此处粘贴，并加盖公章）\n\n\n"),
             size=10, color=RGBColor(0x80, 0x80, 0x80))
    trPr = tbl.rows[0]._tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight'); trHeight.set(qn('w:val'), '3500'); trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)

    doc.add_paragraph()
    doc.add_paragraph()
    h(doc, "二、申请主体为自然人的")
    para(doc, "请提交身份证正反面复印件，并由本人签字确认：")
    tbl2 = doc.add_table(rows=1, cols=2)
    tbl2.style = 'Table Grid'
    for i, label in enumerate(["身份证正面复印件粘贴处", "身份证反面复印件粘贴处"]):
        c = tbl2.rows[0].cells[i]
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(f"\n\n\n（{label}）\n\n\n"), size=10, color=RGBColor(0x80, 0x80, 0x80))
    trPr = tbl2.rows[0]._tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight'); trHeight.set(qn('w:val'), '3500'); trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)

    doc.add_paragraph()
    signature_block(doc)
    doc.save(f)
    print("OK:", f)


def make_authorization_letter():
    """委托代理委托书"""
    f = os.path.join(OUT_DIR, "附属材料4_代理委托书模板.docx")
    doc = Document(); set_section(doc.sections[0])

    title(doc, "计算机软件著作权登记代理委托书")
    doc.add_paragraph()

    para(doc, "委托人（申请人）：____________________________________")
    para(doc, "统一社会信用代码/身份证号：__________________________")
    para(doc, "通讯地址：____________________________________________")
    para(doc, "联系人：__________________  联系电话：________________")
    doc.add_paragraph()
    para(doc, "受托人（代理机构）：__________________________________")
    para(doc, "统一社会信用代码：____________________________________")
    para(doc, "代理人姓名：______________  执业证号/工号：____________")
    para(doc, "联系电话：______________  电子邮箱：__________________")
    doc.add_paragraph()

    para(doc, "委托人现委托受托人作为我方合法代理人，代为办理下列软件著作权登记相关事宜：")
    items = [
        "代为提交“DMS经销商管理系统 V1.0”计算机软件著作权登记申请；",
        "代为接收中国版权保护中心发出的补正、受理、登记证书等通知；",
        "代为办理申请文件补正、陈述意见、撤回申请、领取登记证书等手续；",
        "代为缴纳登记相关费用。",
    ]
    for i, t in enumerate(items, 1):
        para(doc, f"{i}. {t}")
    doc.add_paragraph()
    para(doc,
         "委托期限：自本委托书签署之日起至上述事项办理完毕之日止。"
         "受托人在上述权限范围内所作的行为及签署的文件，委托人均予承认并承担相应法律责任。")
    doc.add_paragraph()
    para(doc, "委托人（盖章/签字）：____________________")
    para(doc, "受托人（盖章）：________________________")
    para(doc, "日　期：　　　 年　　 月　　 日")
    doc.save(f)
    print("OK:", f)


def make_checklist():
    """提交材料清单 & 注意事项"""
    f = os.path.join(OUT_DIR, "附属材料5_提交材料清单与注意事项.docx")
    doc = Document(); set_section(doc.sections[0])

    title(doc, "软件著作权登记提交材料清单与注意事项")
    doc.add_paragraph()

    h(doc, "一、提交材料清单")

    # 表格
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = 'Table Grid'
    headers = ["序号", "材料名称", "份数", "纸质/电子", "备注"]
    for i, hd in enumerate(headers):
        c = tbl.rows[0].cells[i]
        c.text = ""
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(hd), name=u'黑体', size=10, bold=True)
        tcPr = c._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), 'D9E1F2'); tcPr.append(shd)
    rows = [
        ["1", "计算机软件著作权登记申请表", "1", "纸质+电子",
         "在线填写并打印，申请人签字/盖章"],
        ["2", "源程序鉴别材料", "1", "纸质+电子",
         "前30页+后30页，每页50行，页眉写明软件全称版本号"],
        ["3", "文档鉴别材料（用户操作手册）", "1", "纸质+电子",
         "不少于60页，图文结合，页眉写明软件全称版本号"],
        ["4", "申请人主体资格证明", "1", "纸质复印件",
         "企业营业执照副本复印件加盖公章；个人身份证复印件签字"],
        ["5", "经办人身份证复印件", "1", "纸质复印件", "正反面复印，本人签字"],
        ["6", "原创性与不侵权声明", "1", "纸质原件", "申请人签字/盖章"],
        ["7", "代理委托书（如委托代理）", "1", "纸质原件", "委托代理机构时提交"],
        ["8", "合作开发协议（如合作开发）", "1", "纸质复印件", "单独开发无需提供"],
    ]
    for r in rows:
        cs = tbl.add_row().cells
        for i, v in enumerate(r):
            cs[i].text = ""
            p = cs[i].paragraphs[0]
            set_font(p.add_run(v), size=9.5)

    doc.add_paragraph()
    h(doc, "二、文档鉴别材料（源程序、用户手册）格式要求")
    rules = [
        "源程序与文档均使用 A4 纸单面打印，纵向；上、下边距各不少于 2cm，左、右边距各不少于 2cm。",
        "页眉统一标注“DMS经销商管理系统V1.0 源程序”或“DMS经销商管理系统V1.0 用户操作手册”。",
        "每页代码不少于 50 行；空行、注释行可计入 50 行，但注释比例不宜过高。",
        "源程序前 30 页从主程序入口或核心模块开始，后 30 页取代码末尾；如总代码不足 60 页，需全部提交。",
        "中段代码较长时可省略，在第 30 页末尾或第 31 页开头用一行文字注明“中段代码省略，完整代码见随附电子材料”。",
        "用户操作手册不少于 60 页，内容须包含功能说明、操作步骤、界面截图；截图与文字说明一一对应。",
        "源程序与文档中不得出现第三方版权信息、开源协议原文、自动生成标识（如 Generated by、@author AI、Do Not Copy）。",
        "所有提交材料中的软件名称、版本号、开发完成日期必须一致。",
    ]
    for i, t in enumerate(rules, 1):
        para(doc, f"{i}. {t}")

    h(doc, "三、填写注意事项")
    fill_rules = [
        "软件全称一般遵循“企业/品牌 + 软件用途 + 软件”格式，如“DMS 经销商管理系统 V1.0”。",
        "版本号如非 V1.0，须额外提交《版本说明》，说明是原创版本还是升级版本及升级内容。",
        "开发完成日期指本版本整体功能开发完成、通过测试的日期；必须早于首次发表日期。",
        "首次发表状态：若软件尚未向公众提供，选择“未发表”；如已对外使用/销售，填写首次发表日期与地点。",
        "开发方式：单独开发、合作开发、委托开发须如实选择；合作开发须提交合作开发协议；委托开发须提交委托协议。",
        "权利取得方式：原始取得（自己开发）或继受取得（受让/继承/承受）；继受取得须提交相应证明文件。",
        "硬件环境、软件环境、编程语言、源程序量按实际填写，源程序量以总行数计。",
        "申请人名称、证件号码、地址须与证件完全一致，请勿使用简称。",
    ]
    for i, t in enumerate(fill_rules, 1):
        para(doc, f"{i}. {t}")

    h(doc, "四、审核常见补正原因")
    pitfalls = [
        "页眉与申请软件名称/版本号不一致。",
        "源程序每页不足 50 行；最后一页出现大段空白。",
        "源程序中出现第三方公司版权声明（如 Apache、Google、阿里巴巴等）未删除。",
        "用户手册页数不足 60 页或全为文字、没有任何界面截图。",
        "截图与所申请软件名称不一致（例如测试环境里写了其他系统名）。",
        "开发完成日期晚于首次发表日期，或日期晚于申请日。",
        "申请人名称与营业执照不符，或营业执照未加盖公章。",
        "源代码中出现 SQL 脚本、配置文件、package 声明等占比过高，核心业务代码占比不足。",
    ]
    for i, t in enumerate(pitfalls, 1):
        para(doc, f"{i}. {t}")

    h(doc, "五、提交渠道")
    para(doc,
         "中国版权保护中心已全面实行线上登记。请登录“中国版权保护中心”官网（www.ccopyright.com.cn）"
         "注册账号后在线填报申请表，并按系统提示上传上述材料的 PDF 电子版；"
         "审核通过后按通知邮寄纸质材料或全程电子化办理。登记周期一般为 30—60 个工作日。")

    doc.save(f)
    print("OK:", f)


if __name__ == "__main__":
    make_function_intro()
    make_copyright_affidavit()
    make_identity_template()
    make_authorization_letter()
    make_checklist()
